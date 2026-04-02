import os
import io
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def number_to_words(n):
    # A tiny simplified number to words for AED
    try:
        n = int(float(n))
        return f"{n} AED only" # Fallback if we don't import a full library like num2words
    except:
        return f"{n} AED"

def set_cell_background(cell, fill_color):
    """Set background color of a cell"""
    properties = cell._element.tcPr
    if properties is None:
        return
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_color)
    properties.append(shading)

def generate_commercial_proposal_docx(parsed_json, tracking_id):
    """
    Generates a professional Commercial Proposal docx from parsed JSON data.
    """
    doc = Document()
    
    # Optional: adjust margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # 1. Header Table (Invisible borders) for Logo & Title
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run_logo = p_left.add_run("HAMDAZ\n")
    run_logo.font.size = Pt(24)
    run_logo.font.bold = True
    run_logo.font.color.rgb = RGBColor(0, 176, 240) # Hamdaz Blue/Cyan

    p_address = cell_left.add_paragraph()
    p_address.add_run("HAMDAZTECH TECHNOLOGY SERVICES L.L.C\n").font.bold = True
    p_address.add_run("PO Box : 5768, Office 22\nAbu Dhabi, U.A.E\nEmail: hello@hamdaz.com").font.size = Pt(9)
    
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_right.add_run("COMMERCIAL PROPOSAL\n")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    p_right.add_run(f"# QT-{tracking_id[:8].upper()}").font.size = Pt(10)

    doc.add_paragraph() # spacing

    # 2. Bill To & Dates Table
    date_str = datetime.datetime.now().strftime("%d %b %Y")
    expiry_str = (datetime.datetime.now() + datetime.timedelta(days=30)).strftime("%d %b %Y")
    
    info_table = doc.add_table(rows=1, cols=2)
    info_left = info_table.cell(0, 0)
    info_left.paragraphs[0].add_run("Bill To:\n").font.bold = True
    bill_to_name = parsed_json.get("bill_to") or "Customer"
    info_left.add_paragraph(bill_to_name)

    info_right = info_table.cell(0, 1)
    p_dates = info_right.paragraphs[0]
    p_dates.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dates.add_run(f"Quote Date :    {date_str}\n")
    p_dates.add_run(f"Expiry Date :    {expiry_str}")

    doc.add_paragraph()

    # 3. Items Table
    items = parsed_json.get("items", [])
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Table Header
    hdr_cells = table.rows[0].cells
    headers = ['#', 'Item & Description', 'Qty', 'Rate', 'Taxable Amount', 'Tax', 'Amount']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "8E44AD") # Purple theme

    # Add Data
    subtotal = 0.0
    total_tax = 0.0
    for i, item in enumerate(items):
        row_cells = table.add_row().cells
        desc = item.get("description", "Item")
        qty = float(item.get("qty", 1))
        rate = float(item.get("rate", 0))
        tax_rate = float(item.get("tax_rate", 0.05))
        
        taxable_amt = qty * rate
        tax_amt = taxable_amt * tax_rate
        net_amt = taxable_amt + tax_amt
        
        subtotal += taxable_amt
        total_tax += tax_amt
        
        row_cells[0].text = str(i+1)
        row_cells[1].text = desc
        row_cells[2].text = f"{qty:,.2f}"
        row_cells[3].text = f"{rate:,.2f}"
        row_cells[4].text = f"{taxable_amt:,.2f}"
        row_cells[5].text = f"{tax_amt:,.2f}\n( {(tax_rate*100):.1f}% )"
        row_cells[6].text = f"{net_amt:,.2f}"

    doc.add_paragraph()

    # 4. Totals Block
    total_net = subtotal + total_tax
    totals_p = doc.add_paragraph()
    totals_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_p.add_run(f"Sub Total:   {subtotal:,.2f} AED\n").font.bold = True
    totals_p.add_run(f"Total Taxable Amount:   {subtotal:,.2f} AED\n")
    run_tot = totals_p.add_run(f"Total:   {total_net:,.2f} AED\n")
    run_tot.font.bold = True
    run_tot.font.size = Pt(12)
    totals_p.add_run(f"Total in Words:   {number_to_words(total_net)}")

    doc.add_paragraph()

    # 5. Tax Summary
    doc.add_heading('Tax Summary', level=3)
    tax_table = doc.add_table(rows=2, cols=4)
    tax_table.style = 'Table Grid'
    t_hdr = tax_table.rows[0].cells
    for idx, t_text in enumerate(['Tax Details', 'Taxable Amount (AED)', 'Tax Amount (AED)', 'Total Amount (AED)']):
        t_hdr[idx].text = t_text
        t_hdr[idx].paragraphs[0].runs[0].font.bold = True
        set_cell_background(t_hdr[idx], "8E44AD")
        t_hdr[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    t_row = tax_table.rows[1].cells
    t_row[0].text = "Standard Rate (5%)"
    t_row[1].text = f"{subtotal:,.2f}"
    t_row[2].text = f"{total_tax:,.2f}"
    t_row[3].text = f"{total_net:,.2f}"

    doc.add_paragraph()

    # 6. Notes & Terms
    notes = parsed_json.get("notes", "")
    p_notes = doc.add_paragraph()
    p_notes.add_run("Notes:\n").font.bold = True
    p_notes.add_run(notes + "\n\n")
    
    p_notes.add_run("Bank Details:\n").font.bold = True
    p_notes.add_run("Bank Name: Abu Dhabi Commercial Bank\nAccount Title: HAMDAZTECH TECHNOLOGY\nAccount: 10459345934\n\n")

    p_notes.add_run("Terms & Conditions:\n").font.bold = True
    p_notes.add_run("1. All sales shall be under UAE Law.\n2. Payment Terms: As per agreed credit.\n3. Supply details included in delivery note.")

    # Save to memory instead of writing to physical disk
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
