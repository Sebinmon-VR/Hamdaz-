import email
from flask import Flask, redirect, url_for, session, request, render_template, jsonify ,abort ,send_file
from msal import ConfidentialClientApplication
import requests
import os
from dotenv import load_dotenv
import msal
from datetime import datetime, timedelta
import threading
import time
import pandas as pd  # Required for timestamp conversion
from sharepoint_data import *
from sharepoint_items import *
from zoho import *
import openai  
import re
import json
import html
from pinecone import Pinecone, ServerlessSpec
from uuid import uuid4
# from openai import OpenAI
from assistant import *
import docx
import pypdf
import io
from cosmos import (
    get_user_sessions, get_session_messages, save_session_message, delete_session, 
    search_item_distributors, create_shared_project, invite_collaborator, 
    accept_collaboration_invite, get_shared_projects_for_user, get_shared_project_details,
    save_shared_session_message, get_shared_project_activity, update_project_heartbeat,
    get_project_presence, save_user_notification, get_user_notifications, mark_notification_read,
    save_tracked_email, get_tracked_emails_for_task, update_tracked_email_reply, get_pending_tracked_emails
)
# ================== LOAD ENVIRONMENT ==================
load_dotenv(override=True)
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "supersecretkey123")
# ---------------- Azure AD Config ----------------
CLIENT_ID = os.getenv("CLIENT_ID")  
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
REDIRECT_URI = os.getenv("REDIRECT_URI")
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
SCOPE = ["User.Read", "Mail.Send"]
SUPERUSERS = ["jishad@hamdaz.com", "hisham@hamdaz.com" , "sebin@hamdaz.com" ,"sujeel@hamdaz.com","shibit@hamdaz.com", "althaf@hamdaz.com" , "ashna@hamdaz.com"]
approvers = ["shibit@hamdaz.com", "althaf@hamdaz.com" ,"sebin@hamdaz.com" , "sujeel@hamdaz.com", "ashna@hamdaz.com"]
LIMITED_USERS = [""]
# Initialize MSAL
msal_app = ConfidentialClientApplication(
    CLIENT_ID,
    authority=AUTHORITY,
    client_credential=CLIENT_SECRET
)
GRAPH_API_ENDPOINT = "https://graph.microsoft.com/v1.0"
# ---------------- SharePoint Config ----------------
SITE_DOMAIN = "hamdaz1.sharepoint.com"
SITE_PATH = "/sites/ProposalTeam"
LIST_NAME = "Proposals"
test_path = "/sites/Test"
test_proposals_list = "testproposals"
EXCLUDED_USERS = excludeusers_from_sl() 
# ✅ Initialize the OpenAI Client properly
openai.api_key = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")  
pc = Pinecone(api_key=PINECONE_API_KEY)
# Initialize the index
index = pc.Index("hamdaz")
# ==============================================================
# Initialize global data (first load)
print("[INIT] Fetching initial SharePoint data...")
tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
df = items_to_dataframe(tasks)
user_analytics = generate_user_analytics(df, exclude_users=EXCLUDED_USERS)
print("[INIT] Data loaded successfully.")
# ==============================================================
# HELPER FUNCTIONS
# ==============================================================
def is_admin(email_or_name):
    if not email_or_name:
        return False
    identifier_lower = email_or_name.lower().replace(" ", "")
    for superuser in SUPERUSERS:
        # Check if it matches exactly the email, or just the portion before the @
        superuser_email = superuser.lower()
        superuser_name = superuser_email.split('@')[0].replace(" ", "")
        # If the provided name is part of the email prefix, consider it a match
        # e.g., if identifier_lower is "jishad", it matches "jishad@hamdaz.com"
        if identifier_lower == superuser_email or identifier_lower == superuser_name or identifier_lower in superuser_name:
            return True
    return False
def is_approver(email):
    return email.lower() in approvers if email else False
app.jinja_env.globals.update(is_admin=is_admin, is_approver=is_approver, current_date=datetime.now())
def greetings():
    now = datetime.now()
    hour = now.hour
    if 5 <= hour < 12:
        return "Good Morning"
    elif 12 <= hour < 17:
        return "Good Afternoon"
    elif 17 <= hour < 21:
        return "Good Evening"
    else:
        return "Hello"
def get_analytics_data(df, period_type='month', year=None, month=None):
    if year is None:
        year = datetime.now().year
    if month is None:
        month = datetime.now().month
    period = {
        'type': period_type,
        'year': year,
        'month': month
    } if period_type != 'all' else None
    analytics = compute_overall_analytics(df, period)
    per_user = compute_user_analytics_with_last_date(df, EXCLUDED_USERS, period)
    return analytics, per_user
# ==============================================================
# LEAVE EXPIRY HELPER — called by background_updater
# ==============================================================
def check_and_process_expired_leaves():
    """
    Checks all active leave records in Cosmos DB.
    If a leave s end date has passed today, it:
    1. Removes the user from SP excludeusers list
    2. Marks the leave as completed in Cosmos DB
    """
    try:
        from cosmos import get_active_leaves, update_leave_status
        today = datetime.now().date()
        active_leaves = get_active_leaves()
        for leave in active_leaves:
            leave_end_str = leave.get("leave_end", "")
            if not leave_end_str:
                continue
            try:
                leave_end_date = datetime.strptime(leave_end_str, "%Y-%m-%d").date()
            except ValueError:
                continue
            if today > leave_end_date:
                username = leave.get("username", "")
                user_email = leave.get("user_email", "")
                doc_id = leave.get("id", "")
                continue_assign = leave.get("continue_assign", False)
                print(f"[LEAVE-EXPIRY] Leave expired for {username} (end: {leave_end_str})")
                # Remove from excludeusers only if they were added in the first place
                if not continue_assign and username:
                    remove_user_from_excludelist(username)
                # Mark leave as completed in Cosmos
                if doc_id and user_email:
                    update_leave_status(doc_id, user_email, "completed")
    except Exception as e:
        print(f"[LEAVE-EXPIRY ERROR] {e}")
# ==============================================================
# BACKGROUND DATA UPDATER
# ==============================================================
def background_updater():
    """Runs in background to refresh SharePoint data periodically."""
    global tasks, df, user_analytics, EXCLUDED_USERS
    while True:
        try:
            print("[BG] Updating SharePoint data...")
            # Fetch latest SharePoint list
            tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
            
            EXCLUDED_USERS = excludeusers_from_sl()
            
            # list_columns= get_list_columns(SITE_DOMAIN , SITE_PATH , LIST_NAME)
            df = items_to_dataframe(tasks)
            user_analytics = generate_user_analytics(df, exclude_users=EXCLUDED_USERS)
            # Calculate priority score and rank
            user_analytics = calculate_priority_score(user_analytics)
            user_analytics = assign_priority_rank(user_analytics)
            # Fetch existing SharePoint items
            existing_items = get_existing_useranalytics_items()
            # jobcount= user_with_jobs_ls() ////
            for _, row in user_analytics.iterrows():
                username = row["User"]
                item_fields = {
                    "Username": username,
                    "ActiveTasks": int(row["OngoingTasksCount"]),
                    "RecentDate": row["LastAssignedDate"].isoformat() if isinstance(row["LastAssignedDate"], datetime) else row["LastAssignedDate"],
                    "Priority": int(row["PriorityRank"]),
                    # "jobcount": int(jobcount.get(username, 0)) ////
                }
                # ✅ Use safe helper to find existing user
                existing_item = find_existing_user_item(existing_items, username)
                if existing_item:
                    update_user_analytics_in_sharepoint(existing_item["id"], item_fields)
                    print(f"Updated {username} in SharePoint")
                else:
                    add_item_to_sharepoint(item_fields)  
                    print(f"Added new user {username} to SharePoint")
            # Perform smart rotation
            swp()
            # Sync supplier emails in background
            try:
                sync_all_pending_supplier_emails()
            except Exception as se:
                print(f"[BG-SYNC ERROR] {se}")
                
            # Process expired leaves
            try:
                check_and_process_expired_leaves()
            except Exception as le:
                print(f"[BG-LEAVE ERROR] {le}")
                
            print(f"[BG] Data updated successfully at {datetime.now()}", flush=True)
        except Exception as e:
            print("[BG] Error during update:", e)
        time.sleep(100)
# ==============================================================
# ROUTES
# ==============================================================
@app.route("/update_analytics")
def update_analytics():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    period_type = request.args.get('period', 'month')
    year = int(request.args.get('year', datetime.now().year))
    month = int(request.args.get('month', datetime.now().month))
    tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
    df = items_to_dataframe(tasks)
    analytics, per_user = get_analytics_data(df, period_type, year, month)
    return jsonify({
        "analytics": analytics,
        "per_user": per_user
    })
@app.route("/")
def index():
    if "user" in session:
        user = session["user"]
        email = user.get("mail") or user.get("userPrincipalName")
        user_id = user.get("id")
        user_flag_data = get_user_details_from_excell()
        current_user = next((u for u in user_flag_data if u.get("email", "").lower() == email.lower()), None)
        flag_value = current_user.get("flag") if current_user else 0
        try:
            flag = int(flag_value) if str(flag_value).strip() else 0
        except (ValueError, TypeError):
            flag = 0
        if not current_user or flag != 1:
            return redirect(url_for("user_form"))
        if email.lower() in SUPERUSERS:
            dashboard_role = "admin_dashboard"
            excel_role = "admin"
            app.jinja_env.globals.update(excel_role="admin")
        else:
            excel_role = current_user.get("role", "").strip().lower() if current_user else ""
            if excel_role == "pre-sales":
                app.jinja_env.globals.update(excel_role="pre-sales")
                dashboard_role = "pre_sales_dashboard"
            elif excel_role == "business development":
                app.jinja_env.globals.update(excel_role="bd")
                dashboard_role = "business_dev_dashboard"
            elif excel_role == "customer success":
                app.jinja_env.globals.update(excel_role="cs")
                dashboard_role = "customer_success_dashboard"
            elif excel_role == "ai":
                dashboard_role = "admin_dashboard"
                app.jinja_env.globals.update(excel_role="ai")
            else:
                dashboard_role = "user_dashboard"
        period_type = request.args.get('period', 'month')
        year = int(request.args.get('year') or datetime.now().year)
        month = int(request.args.get('month') or datetime.now().month)
        greeting = greetings()
        analytics, per_user = get_analytics_data(df, period_type, year, month)
        username = user.get("displayName", "").replace(" ", "")
        user_analytics_specific = get_user_analytics_specific(df, username)
        now_utc = pd.Timestamp.utcnow()
        ongoing_filtered = [
            t for t in user_analytics_specific['OngoingTasks']
            if pd.to_datetime(t['BCD']) > now_utc and t.get('SubmissionStatus','') != 'Submitted'
        ]
        user_analytics_specific['OngoingTasks'] = ongoing_filtered
        ongoing_tasks_count = len(ongoing_filtered)
        if 'Created' in df.columns:
            df['Created'] = pd.to_datetime(df['Created'])
            available_years = sorted(df['Created'].dt.year.unique().tolist())
        else:
            available_years = [datetime.now().year]
        return render_template(
            f"{dashboard_role}.html",
            role=excel_role,
            user=user,
            greeting=greeting,
            tasks=tasks,
            analytics=analytics,
            per_user=per_user,
            current_period=period_type,
            current_year=year,
            current_month=month,
            available_years=available_years,
            user_analytics=user_analytics_specific,
            email=email,
            ongoing_tasks_count=ongoing_tasks_count,
            due_today_tasks_count=len([
                t for t in user_analytics_specific['OngoingTasks']
                if pd.to_datetime(t['BCD']).date() == now_utc.date()
            ]),
            user_flag_data=user_flag_data, 
        )
    return render_template("login.html")
@app.route("/user_form", methods=["GET", "POST"])
def user_form():
    if "user" not in session:
        return redirect(url_for("login"))
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    user_id = user.get("id")
    if request.method == "POST":
        role = request.form.get("role")
        success = add_or_update_user_in_excel(email, user_id, user.get("displayName"), role)
        if success:
            return redirect(url_for("index"))
        else:
            return "Error saving user data. Please try again."
    return render_template("user_form.html", user=user)
@app.route("/dashboard")
def dashboard():
    return redirect("/")
@app.route("/teams")
def teams():
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    return render_template("teams.html", user=user, email=email, user_analytics=user_analytics)
from urllib.parse import unquote
import time
def get_partnership_data_processed(search="", status="", page=1, per_page=50):
    raw_data = get_partnership_data()# Replace with actual loading!
    filtered = [
        row for row in raw_data
        if (search.lower() in row["Product"].lower() or search.lower() in row["Competitor Company"].lower())
        and (not status or row.get("Status", "Not Started").strip() == status)
    ]
    total_count = len(filtered)
    start = (page - 1) * per_page
    end = start + per_page
    paginated = filtered[start:end]
    return paginated, total_count
from math import ceil
@app.route("/bd")
def view_data():
    user = session.get("user", "BD_USER")  # placeholder session user
    search = request.args.get("search", "")
    status = request.args.get("status", "")
    page = int(request.args.get("page", 1))
    per_page = 50
    data, total_count = get_partnership_data_processed(search, status, page, per_page)
    # Group by Product Group Number
    grouped_data = {}
    for row in data:
        key = row.get("Product Group Number", "N/A")
        grouped_data.setdefault(key, []).append(row)
    total_pages = ceil(total_count / per_page)
    return render_template(
        "business_dev_team.html",
        grouped_data=grouped_data,
        user=user,
        search=search,
        status=status,
        page=page,
        total_pages=total_pages,
        per_page=per_page,
        total_count=total_count
    )
@app.route('/competitor/<path:competitor_name>/<path:product_name>/<path:manufacturer>', methods=['GET', 'POST'])
def competitor_profile(competitor_name, product_name, manufacturer):
    user = session["user"]
    data = get_partnership_data()  # fetch fresh data every time
    # Decode URL and restore slashes
    competitor_name = unquote(competitor_name).strip()
    product_name = unquote(product_name).replace("_slash_", "/").strip()
    manufacturer = unquote(manufacturer).replace("_slash_", "/").strip()
    competitor_entry = next(
        (row for row in data if
         row['Competitor Company'].strip() == competitor_name and
         row['Product'].strip() == product_name and
         row['ADNOC Approved Manufacturer'].strip() == manufacturer),
        None
    )
    if not competitor_entry:
        abort(404)
    other_products = [
        row for row in data
        if row['Competitor Company'].strip() == competitor_name and
        not (row['Product'].strip() == product_name and
             row['ADNOC Approved Manufacturer'].strip() == manufacturer)
    ]
    if request.method == 'POST':
        json_data = request.get_json()
        new_status = json_data.get('status')
        new_remarks = json_data.get('remarks')
        for row in data:
            if row['Competitor Company'].strip() == competitor_name:
                if new_status:
                    save_partnership_update(
                        product_group=row.get('Product Group Number'),
                        product_name=row.get('Product'),
                        manufacturer=row.get('ADNOC Approved Manufacturer'),
                        competitor_name=row.get('Competitor Company'),
                        field='Status',
                        new_value=new_status
                    )
                    row['Status'] = new_status
                if new_remarks:
                    save_partnership_update(
                        product_group=row.get('Product Group Number'),
                        product_name=row.get('Product'),
                        manufacturer=row.get('ADNOC Approved Manufacturer'),
                        competitor_name=row.get('Competitor Company'),
                        field='Remarks',
                        new_value=new_remarks
                    )
                    row['Remarks'] = new_remarks
        return jsonify(success=True)
    return render_template(
        'competitor_profile.html',
        competitor=competitor_entry,
        other_products=other_products,
        user=user
    )
@app.route("/cs")
def cs():
    user = session["user"]
    return render_template("customer_success_team.html", user =user)
@app.route("/user/<username>")
def user_profile(username):
    tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
    df = items_to_dataframe(tasks)
    user_analytics_specific = get_user_analytics_specific(df, username)
    now_utc = pd.Timestamp.utcnow()
    ongoing_filtered = [
        t for t in user_analytics_specific['OngoingTasks']
        if pd.to_datetime(t['BCD']) > now_utc and t.get('SubmissionStatus','') != 'Submitted'
    ]
    user_analytics_specific['OngoingTasks'] = ongoing_filtered
    user_analytics_specific['OngoingTasksCount'] = len(ongoing_filtered)
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    if username in ["dashboard", "customer", "businesscard", "orders", "payments", "reports"]:
        return redirect(f"/{username}")
    return render_template("profile.html", user=user, email=email, user_analytics=user_analytics_specific)
@app.route("/login")
def login():
    auth_url = msal_app.get_authorization_request_url(
        scopes=SCOPE,
        redirect_uri=REDIRECT_URI
    )
    return redirect(auth_url)
@app.route("/getAToken")
def authorized():
    code = request.args.get("code")
    if code:
        result = msal_app.acquire_token_by_authorization_code(
            code,
            scopes=SCOPE,
            redirect_uri=REDIRECT_URI
        )
        if "access_token" in result:
            access_token = result["access_token"]
            graph_data = requests.get(
                "https://graph.microsoft.com/v1.0/me",
                headers={"Authorization": f"Bearer {access_token}"}
            ).json()
            session["user"] = graph_data
            session["access_token"] = access_token
            return redirect("/")
    return "Login failed"
@app.route("/logout")
def logout():
    session.clear()
    return redirect("/")
@app.route("/task_details/<title>")
def task_details(title):
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    df = items_to_dataframe(fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME))
    task = get_task_details(df, title)
    return render_template("pages/task_details.html", task=task, user=user)
@app.route("/businesscard")
def business_cards():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    contacts = get_all_contacts_from_onedrive()
    return render_template("pages/business_cards.html", contacts=contacts, user=user)
@app.route("/api/update-contact", methods=['POST'])
def api_update_contact():
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    data = request.get_json()
    row_id = data.get('row_id')
    contact_data = data.get('contact_data')
    if not row_id or not contact_data:
        return jsonify({"success": False, "error": "Missing data"}), 400
    success = update_contact_in_onedrive_excel(row_id, contact_data)
    if success:
        return jsonify({"success": True})
    else:
        return jsonify({"success": False, "error": "Failed to update Excel file"}), 500
@app.route("/customers")
def customer():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    raw_customers = fetch_customers()
    structured_customers = structure_customers_data(raw_customers)
    return render_template("pages/customers.html", customers=structured_customers, user=user)
@app.route("/quote")
def quote():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    raw_customers = fetch_customers()
    structured_customers = structure_customers_data(raw_customers)
    return render_template("pages/quote.html", user=user, customers=structured_customers)
# ==============================================================
# EMAIL TRACKING AND REPLIES ENDPOINTS
# ==============================================================
@app.route("/api/tracked_emails/<task_id>", methods=["GET"])
def get_tracked_emails(task_id):
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    from cosmos import get_tracked_emails_for_task, task_supplier_quotes_container
    emails = get_tracked_emails_for_task(task_id)
    # Enrich with collection status from task_supplier_quotes
    enriched_emails = []
    for mail in emails:
        status_in_db = None
        if task_supplier_quotes_container:
            query = "SELECT c.status FROM c WHERE c.tracking_id = @tid AND c.task_id = @task_id"
            res = list(task_supplier_quotes_container.query_items(query=query, parameters=[
                {'name':'@tid', 'value': mail['id']},
                {'name':'@task_id', 'value': task_id}
            ], enable_cross_partition_query=True))
            if res:
                status_in_db = res[0].get('status')
        mail['collection_status'] = status_in_db
        enriched_emails.append(mail)
    return jsonify({"success": True, "data": enriched_emails})
def generate_temp_quote_docx(summary, reply_content, tracking_id):
    """Generates a temporary docx based on AI extraction from supplier reply"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()
        doc.add_heading('Supplier Quote Extract', 0)
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        doc.add_heading('Details parsed from email', level=2)
        doc.add_paragraph(reply_content)
        doc.add_paragraph('\n---\nAuto-generated by Hamdaz AI Workspace')
        os.makedirs('static/quotes', exist_ok=True)
        filename = f"static/quotes/Quote_{tracking_id[:8]}.docx"
        doc.save(filename)
        return "/" + filename
    except Exception as e:
        print("Failed to generate DOCX:", e)
        return None
def sync_supplier_emails_for_user(user_email):
    """Refactored to perform targeted searches for each pending email in Cosmos DB."""
    from sharepoint_items import get_access_token
    from cosmos import get_pending_tracked_emails, update_tracked_email_reply, save_session_message, save_task_supplier_quote
    import requests, re, json
    from bs4 import BeautifulSoup
    try:
        app_access_token = get_access_token()
        pending_emails = get_pending_tracked_emails(user_email=user_email)
        if not pending_emails:
            return 0
        headers = {"Authorization": f"Bearer {app_access_token}"}
        matched = 0
        for pe in pending_emails:
            tracking_id = pe['id']
            subject = pe.get('subject', '')
            supplier_email = pe.get('to_email', '')
            task_id = pe['task_id']
            session_id = pe.get('session_id')
            # 1. Perform resilient keyword extraction for search
            # Remove common prefixes/noise to get core keywords
            clean_subject = subject.replace("Re:", "").replace("RE:", "").replace("Fwd:", "").replace("Procurement Inquiry for", "").strip()
            # Split into keywords and take first 3-4 to avoid overly specific queries
            keywords = [k for k in clean_subject.split() if len(k) > 2][:4]
            keyword_search = " ".join(keywords)
            # 2. Perform targeted search in Graph API
            # IMPORTANT: The entire query must be wrapped in double quotes for KQL prefixes like 'from:' to work properly
            search_query = f'from:{supplier_email}'
            if keyword_search:
                search_query += f' {keyword_search}'
            # Wrap the whole thing in double-quotes and then URL encode
            quoted_search = f'"{search_query}"'
            print(f"[SYNC] Searching for {tracking_id} using KQL: {quoted_search}")
            url = f"{GRAPH_API_ENDPOINT}/users/{user_email}/messages?$search={requests.utils.quote(quoted_search)}&$top=10&$select=id,subject,bodyPreview,body,from"
            resp = requests.get(url, headers=headers)
            if resp.status_code != 200:
                print(f"[SYNC] Search failed for {tracking_id}: {resp.text}")
                continue
            messages = resp.json().get("value", [])
            for msg in messages:
                html_body = msg.get("body", {}).get("content", "")
                plainTextPreview = msg.get("bodyPreview", "")
                # Verify match: Either REF ID is present OR Subject/Sender match is strong
                found_ref = re.search(r'REF:([a-f0-9\-]{36})', html_body)
                is_direct_match = (found_ref and found_ref.group(1) == tracking_id)
                # Fallback: if no REF found, check if sender matches (search already filtered subject)
                is_fallback_match = False
                if not found_ref:
                    msg_sender = msg.get('from', {}).get('emailAddress', {}).get('address', '').lower()
                    if msg_sender == supplier_email.lower():
                        is_fallback_match = True
                if is_direct_match or is_fallback_match:
                    try:
                        soup = BeautifulSoup(html_body, "html.parser")
                        clean_text = soup.get_text(separator=' ').strip()
                    except:
                        clean_text = plainTextPreview
                    # Mark read
                    requests.post(f"{GRAPH_API_ENDPOINT}/users/{user_email}/messages/{msg['id']}", headers=headers, json={"isRead": True})
                    # AI Extraction
                    ai_prompt = f"""
                    Analyze this supplier reply for procurement context. Extract pricing/quote details into JSON.
                    {{
                        "is_quote": true,
                        "bill_to": "Recipient name",
                        "items": [{{"description": "...", "qty": 1, "rate": 0}}],
                        "summary": "1-sentence summary of the reply",
                        "notes": "..."
                    }}
                    Reply content: {clean_text[:5000]}
                    """
                    try:
                        from openai import OpenAI
                        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
                        ai_res = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[{"role": "system", "content": "Return raw JSON only."}, {"role": "user", "content": ai_prompt}],
                            temperature=0.1
                        )
                        raw_ai = ai_res.choices[0].message.content.strip().replace("```json","").replace("```","")
                        parsed = json.loads(raw_ai)
                        summary = parsed.get("summary", "New reply received.")
                        quote_doc_path = f'/api/quotes/download/{tracking_id}' if parsed.get("is_quote") else None
                        if parsed.get("is_quote"):
                            save_task_supplier_quote(task_id, tracking_id, supplier_email, summary, parsed)
                        update_tracked_email_reply(tracking_id, task_id, clean_text, summary, quote_doc_path, ai_parsed_data=parsed)
                        if session_id:
                            msg_content = f"**Supplier Reply Received ({supplier_email})!**\n\n{summary}"
                            if quote_doc_path: msg_content += f"\n\n[ðŸ“¥ Download Commercial Proposal]({quote_doc_path})"
                            save_session_message(session_id, user_email, "assistant", msg_content, agent_type="procurement", task_id=task_id)
                        matched += 1
                        break # Found the reply for this specific tracked item
                    except Exception as e:
                        print(f"[SYNC] AI Error for {tracking_id}: {e}")
        return matched
    except Exception as e:
        print(f"[SYNC ERROR] {user_email}: {e}")
        return 0
@app.route("/api/check_email_replies", methods=["POST"])
def check_email_replies():
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user_email = session["user"].get("mail") or session["user"].get("userPrincipalName", "")
    matched = sync_supplier_emails_for_user(user_email)
    return jsonify({"success": True, "matched": matched})
def sync_all_pending_supplier_emails():
    """Global background function to sync emails for all users with pending tracked emails."""
    from cosmos import get_pending_tracked_emails
    try:
        pending = get_pending_tracked_emails(user_email=None) # Get all pending
        if not pending:
            return
        unique_users = {pe['user_email'] for pe in pending if 'user_email' in pe}
        print(f"[BG-SYNC] Checking replies for {len(unique_users)} users: {unique_users}")
        for user_email in unique_users:
            try:
                count = sync_supplier_emails_for_user(user_email)
                if count > 0:
                    print(f"[BG-SYNC] Found {count} new replies for {user_email}")
            except Exception as e:
                print(f"[BG-SYNC ERROR] User {user_email}: {e}")
    except Exception as e:
        print(f"[BG-SYNC ERROR] Global loop: {e}")
# ==============================================================
def get_first(quote_data, key, index=0, default=""):
    """
    Safely get the first or indexed value from a list in quote_data.
    """
    values = quote_data.get(key, [default])
    if isinstance(values, list):
        if len(values) > index:
            return values[index]
        else:
            return default
    return values or default
@app.route("/send_quote_for_approval", methods=["POST"])
def send_for_approval():
    if "user" not in session:
        return redirect(url_for("login"))
    user = session.get("user")
    submitter_email = user.get("mail") or user.get("userPrincipalName")
    quote_data = request.form.to_dict(flat=False)
    supplier_file = request.files.get("supplier_quote")  # Uploaded file
    def safe_float(val, default=0):
        try:
            return float(val)
        except (TypeError, ValueError):
            return default
    # Process quote line items as before
    num_items = len(quote_data.get("item_details[]", []))
    combined_items = []
    total_amount = 0
    total_selling = 0
    total_discount = 0
    total_tax_amount = 0
    total_margin = 0
    total_rate = 0
    count = 0
    for i in range(num_items):
        qty = int(quote_data.get("quantity[]", [])[i] or 0)
        rate = safe_float(quote_data.get("rate[]", [])[i])
        margin = safe_float(quote_data.get("margin[]", [])[i])
        discount = safe_float(quote_data.get("discount[]", [])[i])
        tax = safe_float(quote_data.get("tax[]", [])[i])
        base_amount = qty * rate * (1 + margin / 100)
        amount = base_amount - discount
        tax_amount = amount * tax
        selling = amount + tax_amount
        combined_items.append({
            "ItemDetails": quote_data.get("item_details[]", [])[i],
            "Brand": quote_data.get("brand[]", [])[i],
            "Quantity": qty,
            "Rate": rate,
            "Margin": margin,
            "Tax": tax,
            "Discount": discount,
            "Amount": amount,
            "SellingPrice": selling
        })
        total_amount += amount
        total_selling += selling
        total_discount += discount
        total_tax_amount += tax_amount
        total_margin += margin
        total_rate += rate * qty
        count += 1
    avg_tax_percentage = (total_tax_amount / total_amount * 100) if total_amount else 0
    item_fields = {
        "Title": quote_data.get("reference", ["No Title"])[0],
        "CustomerID": quote_data.get("customer_id", [""])[0],
        "Currency": quote_data.get("currency", [""])[0],
        "PaymentTerms": quote_data.get("payment_terms", [""])[0],
        "Email": quote_data.get("email", [""])[0],
        "TaxTreatment": quote_data.get("tax_treatment", [""])[0],
        "Reference": quote_data.get("reference", [""])[0],
        "QuoteDate": quote_data.get("quote_date", [""])[0],
        "ExpiryDate": quote_data.get("expiry_date", [""])[0],
        "Portal": quote_data.get("portal", [""])[0],
        "QuoteCreator": quote_data.get("quote_creator", [""])[0],
        "BCD": quote_data.get("bcd", [""])[0],
        "ApprovalStatus": "Pending",
        "Tax": avg_tax_percentage,
        "Amount": total_amount,
        "TotalDiscount": total_discount,
        "TotalSellingPrice": total_selling,
        "Margin": (total_margin / count) if count else 0,
        "Rate": (total_rate / sum([int(q) for q in quote_data.get("quantity[]", [])])) if count else 0,
        "AllItems": json.dumps(combined_items, indent=2)
    }
    print("DEBUG SharePoint payload:", json.dumps(item_fields, indent=2))
    try:
        # Add SharePoint list item
        resp_json = add_sharepoint_list_item(item_fields)
        if not resp_json:
            raise Exception("Failed to create SharePoint item")
        item_id = resp_json.get("id")
        if supplier_file and supplier_file.filename:
            print("File detected in request")
            file_name = supplier_file.filename  # Use secure_filename if needed
            file_bytes = supplier_file.read()
            print(f"File size: {len(file_bytes)} bytes")
            # Upload file to SharePoint document library folder (replace 'QuoteAttachments' with your folder path)
            share_link = upload_file_to_sharepoint_folder( folder_path="Shared Documents/QuoteAttachments", file_name=file_name, file_bytes=file_bytes)
            print("File uploaded. Share link:", share_link)
            # Update SharePoint list item with link
            update_sharepoint_item_with_link(item_id, share_link)
        else:
            print("No file in request or filename missing")
        return render_template("pages/quote_success.html", user=user, added_items=1)
    except Exception as e:
        print("SharePoint Error:", e)
        return f"âŒ Error adding quote to SharePoint: {str(e)}", 500
@app.route("/quote_decision")
def quote_decision():
    if "user" not in session:
        return redirect(url_for("login"))
    user = session.get("user")
    user_name = user.get("displayName")  # user's name in session
    try:
        site_domain = "hamdaz1.sharepoint.com"
        site_path = "/sites/Test"
        list_name = "Quotes"
        quote_items = fetch_sharepoint_list(site_domain, site_path, list_name)
        # Filter only quotes created by this user
        quote_items = [q for q in quote_items if q.get("QuoteCreator") == user_name]
    except Exception as e:
        print(f"Error fetching SharePoint list items: {e}")
        quote_items = []
    return render_template("pages/quote_decision.html", user=user, quote_items=quote_items)
@app.route("/quote_details/<quote_id>")
def quote_details(quote_id):
    if "user" not in session:                                   
        return redirect(url_for("login"))
    user = session.get("user")
    # Fetch quote by ID from SharePoint
    site_domain = "hamdaz1.sharepoint.com"
    site_path = "/sites/Test"
    list_name = "Quotes"
    quote_items = fetch_sharepoint_list(site_domain, site_path, list_name)
    quote = next((q for q in quote_items if str(q.get("id")) == str(quote_id)), None)
    if not quote:
        return "Quote not found", 404
    # Get customer name from Zoho
    customer_id = quote.get("CustomerID", "")
    customer_name = get_customer_name_from_zoho(customer_id) or ""
    quote["CustomerName"] = customer_name
    attachment_str = quote.get("attachmentlink", "")
    if attachment_str:
        quote["AttachmentLinks"] = [link.strip() for link in attachment_str.split(",") if link.strip()]
    else:
        quote["AttachmentLinks"] = []
    # Check Approval Status
    approval_status = quote.get("ApprovalStatus", "")
    is_approved = approval_status.lower() == "approved"
    quote["IsApproved"] = is_approved
    if is_approved:
        print("approved")
    # -----------------------------
    # Parse AllItems JSON
    # -----------------------------
    all_items_raw = quote.get("AllItems", "")
    try:
        match = re.search(r'\[.*\]', html.unescape(all_items_raw), re.DOTALL)
        if match:
            items = json.loads(match.group(0))
            for item in items:
                item.setdefault('Discount', 0)
                item.setdefault('Quantity', 1)  # default to 1 if missing
            quote['AllItems_parsed'] = items
            def to_float(val):
                try:
                    if val is None:
                        return 0.0
                    if isinstance(val, str):
                        val = val.replace("%", "").strip()
                    return float(val)
                except:
                    return 0.0
            # -----------------------------
            # Calculate Totals correctly
            # -----------------------------
            total_rate = total_amount = total_selling_price = total_tax = total_discount = 0.0
            total_margin_value = 0.0
            margin_count = 0
            total_quantity = 0.0
            for item in items:
                quantity = to_float(item.get('Quantity', 1))
                rate = to_float(item.get('Rate', 0))
                margin = to_float(item.get('Margin', 0))
                tax = to_float(item.get('Tax', 0))
                discount = to_float(item.get('Discount', 0))
                amount = rate * quantity
                item['Amount'] = amount
                tax_amount = amount * tax / 100.0
                selling_price = amount * (1 + margin / 100.0) + tax_amount - (discount * quantity)
                qty = item['Quantity']
                base_amount = (qty * rate * (1 + margin / 100))
                amount = (base_amount - discount)
                tax_amount = (amount * tax)
                selling = (amount + tax_amount)
                item['SellingPrice'] = selling
                # Accumulate totals per item
                total_selling_price += selling
                total_tax += tax_amount
                total_rate += amount
                total_amount += amount
                total_discount += discount * quantity
                if margin > 0:
                    total_margin_value += margin
                    margin_count += 1
                total_quantity += quantity
            avg_margin = total_margin_value / margin_count if margin_count > 0 else 0
            avg_rate = total_rate / total_quantity if total_quantity > 0 else 0
            quote['Totals'] = {
                "Rate": total_rate,
                "AvgRate": avg_rate,
                "AvgMargin": avg_margin,
                "Tax": total_tax,
                "Amount": total_amount,
                "SellingPrice": total_selling_price,
                "Discount": total_discount
            }
        else:
            quote['AllItems_parsed'] = []
            quote['Totals'] = {}
    except Exception as e:
        print("Error parsing AllItems:", e)
        quote['AllItems_parsed'] = []
        quote['Totals'] = {}
    return render_template("pages/quote_details.html", quote=quote, user=user)
@app.route("/export_quote/<quote_id>")
def export_quote(quote_id):
    if "user" not in session:
        return redirect(url_for("login"))
    # --- REUSE YOUR EXISTING FETCH LOGIC HERE ---
    # (Copy the logic from quote_details to fetch and parse the quote)
    site_domain = "hamdaz1.sharepoint.com"
    site_path = "/sites/Test"
    list_name = "Quotes"
    quote_items = fetch_sharepoint_list(site_domain, site_path, list_name)
    quote = next((q for q in quote_items if str(q.get("id")) == str(quote_id)), None)
    if not quote:
        return "Quote not found", 404
    # Fetch Customer Name
    customer_id = quote.get("CustomerID", "")
    customer_name = get_customer_name_from_zoho(customer_id) or ""
    quote["CustomerName"] = customer_name
    # Parse AllItems (Required for the loop)
    import json, re, html
    all_items_raw = quote.get("AllItems", "")
    try:
        match = re.search(r'\[.*\]', html.unescape(all_items_raw), re.DOTALL)
        if match:
            quote['AllItems_parsed'] = json.loads(match.group(0))
        else:
            quote['AllItems_parsed'] = []
    except:
        quote['AllItems_parsed'] = []
    # ---------------------------------------------
    # Generate the Excel file
    excel_file = generate_quote_excel(quote)
    # Return as a download
    return send_file(
        excel_file,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f"Quote_{quote_id}.xlsx"
    )
@app.route("/vendors")
def vendors():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    vendors = get_excel_data_from_onedrive("Partnership_Status.xlsx", "Vendor-Partnership")
    df=pd.DataFrame(vendors)
    # Check the columns first
    print(df.columns)
    # Example: select only useful columns
    columns_of_interest = ['text', 'values']  # or actual column names in your df
    df_selected = df[columns_of_interest]
    # If 'values' contains the actual row data
    # Convert 'values' column (list) into separate columns
    df_expanded = pd.DataFrame(df['values'].tolist(), columns=[
        'ID', 'Vendor', 'Col3', 'Status', 'Col5', 'Comments', 'URL', 'Admin User', 'Password', 'Col10', 'Contact', 'Col12'
    ])
    # View the cleaned DataFrame
    print(df_expanded.head())
    return render_template("pages/vendors.html", data=df_expanded.to_dict(orient="records"), user=user)
@app.route("/vendor/<vendor_id>")
def vendor_detail(vendor_id):
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    vendors = get_excel_data_from_onedrive("Partnership_Status.xlsx", "Vendor-Partnership")
    df = pd.DataFrame(vendors)
    df_expanded = pd.DataFrame(df['values'].tolist(), columns=[
        'ID', 'Vendor', 'Col3', 'Status', 'Col5', 'Comments', 'URL', 'Admin User', 'Password', 'Col10', 'Contact', 'Col12'
    ])
    # Convert ID to string for comparison
    vendor = df_expanded[df_expanded['ID'].astype(str) == str(vendor_id)].to_dict(orient="records")
    if not vendor:
        return "Vendor not found", 404
    return render_template("pages/vendor_detail.html", vendor=vendor[0], user=user)
# ==============================================================
@app.route("/approvals")
def approvals():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    site_domain = "hamdaz1.sharepoint.com"
    site_path = "/sites/Test"
    list_name = "Quotes"
    quote_items = fetch_sharepoint_list(site_domain, site_path, list_name)
    # Show all the quotes with ApprovalStatus as both 'Pending' and 'Approved'
    quote_items = [q for q in quote_items if q.get("ApprovalStatus") in ["Pending", "Approved"]]
    return render_template("pages/quote_decision.html", user=user, quote_items=quote_items)
# ==============================================================
# ==============================================================
pa_chat_histories = {} 
# Note: pa_chat_histories is now legacy, using Cosmos DB instead
@app.route("/personal_assistant")
def personal_assistant_page():
    if "user" not in session:
        return redirect(url_for("login"))
    user = session["user"]
    return render_template("pages/personal_assistant.html", user=user)
@app.route("/api/personal_assistant/sessions", methods=["GET"])
def pa_get_sessions():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    sessions = get_user_sessions(email)
    return jsonify({"sessions": sessions})
@app.route("/api/personal_assistant/sessions/<session_id>", methods=["GET"])
def pa_get_session(session_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    from cosmos import sessions_container
    try:
        session_doc = sessions_container.read_item(item=session_id, partition_key=session_id)
        return jsonify({
            "messages": session_doc.get("messages", []),
            "agent_type": session_doc.get("agent_type", "personal"),
            "session_title": session_doc.get("session_title", "Chat")
        })
    except Exception as e:
        print(f"Error fetching session: {e}")
        return jsonify({"error": "Session not found"}), 404
@app.route("/api/personal_assistant/log", methods=["POST"])
def pa_save_log():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    data = request.json
    session_id = data.get("session_id")
    message = data.get("message")
    status = data.get("status", "info") # success, error, info
    task_id = data.get("task_id")
    if not session_id or not message:
        return jsonify({"error": "Missing session_id or message"}), 400
    from cosmos import save_session_message
    content = f"[LOG][{status.upper()}] {message}"
    save_session_message(session_id, user_email, "system_log", content, agent_type="procurement", task_id=task_id)
    return jsonify({"success": True})
@app.route("/api/personal_assistant/sessions/<session_id>", methods=["DELETE"])
def pa_delete_session(session_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    success = delete_session(session_id)
    if success:
        return jsonify({"success": True})
    return jsonify({"error": "Delete failed"}), 500
@app.route("/api/personal_assistant/chat", methods=["POST"])
def pa_chat():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    username = user.get("displayName", "").replace(" ", "")
    is_admin_user = is_admin(email)
    message = request.form.get("message", "")
    session_id = request.form.get("session_id", "")
    agent_type = request.form.get("agent_type", "personal")
    files = request.files.getlist("files")
    is_analysis = request.form.get("is_analysis", "false").lower() == "true"
    files_text = ""
    if files:
        for f in files:
            if not f.filename:
                continue
            ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
            file_bytes = f.read()
            if not file_bytes:
                continue
            files_text += f"\n--- {f.filename} ---\n"
            try:
                import PyPDF2
                import io
                if ext == 'pdf':
                    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                    for page in reader.pages:
                        files_text += (page.extract_text() or "") + "\n"
                elif ext in ['xlsx', 'xls']:
                    df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                    for sheet_name, df_sheet in df_dict.items():
                        files_text += f"\n--- Sheet: {sheet_name} ---\n"
                        files_text += df_sheet.to_csv(index=False) + "\n"
                elif ext == 'csv':
                    df_file = pd.read_csv(io.BytesIO(file_bytes))
                    files_text += df_file.to_csv(index=False) + "\n"
                elif ext == 'docx':
                    doc = docx.Document(io.BytesIO(file_bytes))
                    files_text += "\n".join([p.text for p in doc.paragraphs]) + "\n"
                elif ext == 'doc':
                    files_text += "[Error reading file: Legacy .doc format is not supported. Please save as .docx and re-upload.]\n"
                elif ext == 'txt':
                    files_text += file_bytes.decode('utf-8') + "\n"
            except Exception as e:
                files_text += f"[Error reading file: {str(e)}]\n"
    # Get chat history from Cosmos DB
    chat_history = []
    print(f"[PA_CHAT] session_id received: '{session_id}'", flush=True)
    if session_id and session_id not in ("", "null", "undefined"):
        msgs = get_session_messages(session_id)
        if msgs:
            # Clean history for OpenAI: filter out system logs and strip timestamps
            chat_history = [{"role": m["role"], "content": m["content"]} for m in msgs if m.get("role") != "system_log"]
            print(f"[PA_CHAT] Prepared {len(chat_history)} messages for AI (filtered logs)", flush=True)
        else:
            print(f"[PA_CHAT] No messages found for session {session_id}", flush=True)
    else:
        print("[PA_CHAT] No session_id — starting fresh conversation", flush=True)
    # Save user message to Cosmos
    title = message[:40].strip() + ("..." if len(message) > 40 else "") if message else "Chat with Assistant"
    # Only set title on the FIRST message (when session_id is empty — new session)
    first_message = not session_id or session_id in ("", "null", "undefined")
    if not first_message:
        title = None  # Don't override existing session title
    print(f"[PA_CHAT] Saving user message. title={title}, first_message={first_message}, agent={agent_type}", flush=True)
    generated_session_id = save_session_message(session_id, email, "user", message, title=title, agent_type=agent_type)
    print(f"[PA_CHAT] Generated/Updated session_id: {generated_session_id}", flush=True)
    try:
        system_instr = ""
        if is_analysis:
            system_instr = "User has likely uploaded a procurement document. ANALYZE it and return the required items STRICTLY as a JSON list of objects with 'name' and 'type' properties. DO NOT add any conversational text before or after the JSON array."
        # If we have system instructions, inject them into the run_personal_assistant call if possible
        # or just prepend to the message if run_personal_assistant doesn't support them.
        # Looking at run_personal_assistant signature in assistant.py...
        reply = run_personal_assistant(username, message, files_text, chat_history, is_admin_user, system_instr=system_instr)
        # Save assistant reply to Cosmos (ensure agent_type is passed)
        save_session_message(generated_session_id, email, "assistant", reply, agent_type=agent_type)
        return jsonify({"reply": reply, "session_id": generated_session_id, "agent_type": agent_type})
    except Exception as e:
        import traceback
        return jsonify({"error": str(e) + "\n" + traceback.format_exc()}), 500
@app.route("/updates")
def updates():
    if "user" not in session:
        return redirect(url_for("login"))
    return render_template("pages/updates.html", user=session["user"])
# ==============================================================
# MS GRAPH SEND EMAIL ROUTE
# ==============================================================
@app.route("/api/send_email", methods=["POST"])
def send_email_api():
    if "user" not in session or "access_token" not in session:
        return jsonify({"success": False, "error": "Unauthorized. Please log in again."}), 401
    try:
        data = request.get_json()
        to_email = data.get("to")
        subject = data.get("subject")
        body = data.get("body")
        task_id = data.get("task_id") # new field for tracking
        session_id = data.get("session_id") # new field for chat session injection
        if not to_email or not subject or not body:
            return jsonify({"success": False, "error": "Missing 'to', 'subject', or 'body' fields."}), 400
        user_email = session["user"].get("mail") or session["user"].get("userPrincipalName", "")
        tracking_id = str(uuid4())
        # If a task ID is provided, we should track this email
        parsed_body = body
        if task_id:
            parsed_body += f"<br><br><span style='color: transparent; font-size: 0px;'>REF:{tracking_id}</span>"
        # Use the session access token directly (must have Mail.Send scope)
        # This works after admin grants Mail.Send consent in Azure Portal
        # and user re-logs in to get a fresh token with Mail.Send.
        access_token = session.get("access_token")
        if not access_token:
            return jsonify({"success": False, "error": "Session expired. Please log in again."}), 401
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        # Handle multiple recipients (comma-separated or list)
        if isinstance(to_email, str):
            to_recipients = [{"emailAddress": {"address": email.strip()}} for email in to_email.split(",")]
        else:
            to_recipients = [{"emailAddress": {"address": email.strip()}} for email in to_email]
        email_data = {
            "message": {
                "subject": subject,
                "body": {
                    "contentType": "HTML", # Changed to HTML to hide tracking ID
                    "content": parsed_body
                },
                "toRecipients": to_recipients
            },
            "saveToSentItems": "true"
        }
        response = requests.post(f"{GRAPH_API_ENDPOINT}/me/sendMail", headers=headers, json=email_data)
        if response.status_code == 202:
            if task_id:
                save_tracked_email(task_id, session_id, to_email, subject, tracking_id, user_email, body)
            return jsonify({"success": True, "message": "Email sent successfully."})
        else:
            error_detail = ""
            try:
                error_detail = response.json().get("error", {}).get("message", response.text)
            except Exception:
                error_detail = response.text
            print(f"[GRAPH EMAIL ERROR] Status: {response.status_code} | Detail: {error_detail}")
            return jsonify({"success": False, "error": f"Graph API Error ({response.status_code}): {error_detail}"}), response.status_code
    except Exception as e:
        import traceback
        return jsonify({"success": False, "error": str(e) + "\n" + traceback.format_exc()}), 500
# ==============================================================
# ==============================================================
@app.route("/user_report")
def user_report():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    user_name = user.get("displayName").replace(" ", "")
    tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
    df = items_to_dataframe(tasks)
    user_analytics_specific = get_user_analytics_specific(df, user_name)
    return render_template("pages/user_report.html", user=user, user_analytics=user_analytics_specific)
@app.route("/admin_report")
def admin_report():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session.get("user")
    tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
    df = items_to_dataframe(tasks)
    overall_analytics, per_user_analytics = get_analytics_data(df, period_type='all')
    return render_template("pages/admin_report.html", user=user, overall_analytics=overall_analytics, per_user_analytics=per_user_analytics)
# ==============================================================
from PyPDF2 import PdfMerger
import io
DEFAULT_PDF_PATH = 'default.pdf'
@app.route('/merge', methods=['GET', 'POST'])
def merge():
    if request.method == 'POST':
        # Get PDFs from the form
        uploaded_files = request.files.getlist('pdf_files')
        # Prepare in-memory PDFs
        pdf_streams = []
        for f in uploaded_files:
            pdf_streams.append(io.BytesIO(f.read()))
        # Include default PDF if requested
        include_default = request.form.get('include_default')
        if include_default == 'on':
            default_pdf_position = int(request.form.get('default_position', 0))
            with open(DEFAULT_PDF_PATH, 'rb') as df:
                default_pdf_stream = io.BytesIO(df.read())
            pdf_streams.insert(default_pdf_position, default_pdf_stream)
        # Merge PDFs
        merger = PdfMerger()
        for pdf_io in pdf_streams:
            merger.append(pdf_io)
        output_pdf = io.BytesIO()
        merger.write(output_pdf)
        merger.close()
        output_pdf.seek(0)
        return send_file(output_pdf, as_attachment=True, download_name='merged.pdf', mimetype='application/pdf')
    return render_template('merge_tp.html' , user=session.get("user"))
import PyPDF2
# Make sure to set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")
CHUNK_SIZE = 3000  # characters per chunk
@app.route("/tp")
def tp():
    try:
        items = get_child_files()
        files = []
        for f in items:
            if "file" in f:
                files.append({
                    "name": f["name"],
                    "url": f["webUrl"],
                    "size": f["size"],
                    "modified": f["lastModifiedDateTime"],
                    "id": f["id"]
                })
        # Default attachments for table
        default_attachments = [
            {"name": f["name"], "id": f["id"], "url": f["url"]}
            for f in files
        ]
    except Exception as e:
        print("Error:", e)
        files = []
        default_attachments = []
    return render_template("tp.html", files=files, default_attachments=default_attachments, user=session.get("user"))
# ddlf
# --- ADD THIS NEW ROUTE ---
@app.route("/download/docx/<file_id>")
def download_source_docx(file_id):
    try:
        # 1. Download the raw DOCX from your source (SharePoint/Graph)
        docx_path = download_docx(file_id)
        # 2. Send the DOCX directly (NO PDF CONVERSION)
        return send_file(
            docx_path, 
            as_attachment=True, 
            download_name=os.path.basename(docx_path),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print("Error downloading DOCX:", e)
        return f"Error: {e}", 500
@app.route("/download/file/<file_id>")
def download_file_from_onedrive(file_id):
    try:
        file_path, file_name = download_file(file_id)
        return send_file(
            file_path,
            as_attachment=True,
            download_name=file_name
        )
    except Exception as e:
        print("Error downloading file:", e)
        return f"Error: {e}", 500
@app.route('/analyze', methods=['POST'])
def analyze_document():
    # --- Input Validation ---
    if 'pdf_file' not in request.files:
        return jsonify({"error": "No PDF file uploaded"}), 400
    pdf_file = request.files['pdf_file']
    # --- Step 1: Extract text from PDF ---
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
    except Exception as e:
        return jsonify({"error": f"Failed to read PDF: {str(e)}"}), 400
    text_content = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            # Clean extra whitespace & preserve structure
            page_text = "\n".join([line.strip() for line in page_text.splitlines() if line.strip()])
            text_content += page_text + "\n"
    if not text_content.strip():
        return jsonify({"error": "No text could be extracted from PDF"}), 400
    # --- Step 2: Split text into manageable chunks ---
    chunks = [text_content[i:i+CHUNK_SIZE] for i in range(0, len(text_content), CHUNK_SIZE)]
    # --- Step 3: Prepare final result (UPDATED) ---
    final_result = {
        "requirements": [],
        "attachments_needed": [],
        "eligibility_criteria": [],
        "deadlines": [],
        "technical_specifications": [],
        "other_notes": [],
        "required_items_needed": [] 
    }
    # --- Step 4: Analyze each chunk with OpenAI ---
    for chunk in chunks:
        prompt = f"""
You are an expert RFQ/SOW analyst.
Step 1: Read the document chunk below.
Step 2: Extract all relevant information. For "required_items_needed", you MUST classify each item as "product", "service", or "note".
Return your answer in JSON exactly like this:
{{
  "requirements": [...],
  "attachments_needed": [...],
  "eligibility_criteria": [...],
  "deadlines": [...],
  "technical_specifications": [...],
  "other_notes": [...],
  "required_items_needed": [
      {{ "name": "Item Name 1", "type": "product" }},
      {{ "name": "A general note about delivery", "type": "note" }}
  ]
}}
If a field is missing, return an empty list.
Document Text:
{chunk}
        """
        try:
            # NOTE: Ensure 'openai' library is configured before running.
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert RFQ/SOW analyst."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0,
                max_tokens=2000
            )
            ai_output = response.choices[0].message.content
            ai_output_clean = ai_output.replace("```json", "").replace("```", "").strip()
            try:
                chunk_data = json.loads(ai_output_clean)
                # Merge chunk data into final result
                # This loop handles all keys, including the new one
                for key in final_result.keys():
                    if key in chunk_data and isinstance(chunk_data[key], list):
                        final_result[key].extend(chunk_data[key])
            except json.JSONDecodeError:
                # fallback for unparsable chunk
                final_result.setdefault("raw_text_chunks", []).append(ai_output_clean)
        except Exception as e:
            # Handle potential API errors (e.g., network, rate limit)
            return jsonify({"error": f"OpenAI API Error: {str(e)}"}), 500
    # --- Step 5: Final Cleanup and Return ---
    # Remove duplicates from lists
    for key in final_result:
        if isinstance(final_result[key], list):
            if key == 'required_items_needed':
                # Handle list of dicts: remove duplicates based on 'name'
                seen = set()
                unique_items = []
                for item in final_result[key]:
                    # Ensure item is a dictionary and has a 'name'
                    if isinstance(item, dict) and 'name' in item:
                        if item['name'] not in seen:
                            unique_items.append(item)
                            seen.add(item['name'])
                    elif isinstance(item, str): # Keep strings if they appear
                        if item not in seen:
                            unique_items.append({"name": item, "type": "product"}) # Convert to new format
                            seen.add(item)
                final_result[key] = unique_items
            else:
                # Using dict.fromkeys to efficiently remove duplicates while preserving order for hashable types
                try:
                    final_result[key] = list(dict.fromkeys(final_result[key]))
                except TypeError:
                    # Fallback for lists with unhashable types that are not 'required_items_needed'
                    # This is a safe fallback, though not expected with current prompt
                    pass
    return jsonify({"extracted_data": final_result})
@app.route('/find_distributors', methods=['POST'])
def find_distributors():
    try:
        data = request.get_json()
        item_name = data.get('item_name')
        location = data.get('location') # Expected to be 'UAE'
        if not item_name:
            return jsonify({"error": "Item name is required"}), 400
        # This query uses the OpenAI model to perform the search/analysis
        prompt = f"""
        Find and list at least 3 authorized distributors or major retailers for '{item_name}' in the {location}.
        For each distributor, provide their official name and a direct link to their company profile or a relevant product page.
        Return your answer in JSON exactly like this:
        {{
          "distributors": [
            {{
              "name": "Distributor Name 1",
              "link": "https://profile.link.1"
            }},
            {{
              "name": "Distributor Name 2",
              "link": "https://profile.link.2"
            }}
          ]
        }}
        If no distributors are found, return an empty list for 'distributors'.
        """
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        response = client.chat.completions.create(
            model="gpt-4o", # Use a model with strong web-browsing capabilities
            messages=[
                {"role": "system", "content": "You are a specialized procurement agent who searches the web for product distributors and returns results in a clean JSON format."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=1000
        )
        ai_output = response.choices[0].message.content
        ai_output_clean = ai_output.replace("```json", "").replace("```", "").strip()
        try:
            # Parse and return the JSON directly to the frontend
            distributor_data = json.loads(ai_output_clean)
            return jsonify(distributor_data), 200
        except json.JSONDecodeError:
            return jsonify({"error": "AI returned unparsable JSON response.", "raw_output": ai_output_clean}), 500
    except Exception as e:
        # Catch API key errors, network errors, etc.
        return jsonify({"error": f"Failed to find distributors: {str(e)}"}), 500
from werkzeug.datastructures import FileStorage
def get_file_extension(filename):
    """Returns the file extension in lowercase."""
    return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
from pypdf import PdfReader # <-- This line is CRITICAL and must be at the top!
@app.route('/rfq_vs_quotes', methods=['GET'])
def load_page():
    """Renders the HTML file uploader page."""
    # Assumes HTML file is at 'templates/pages/rfq_vs_quotes.html'
    return render_template("pages/rfq_vs_quotes.html", user=session.get("user"))
@app.route('/rfq_vs_quotes', methods=['POST'])
def process_files():
    """
    Handles file upload, parses content (PDF, Excel, CSV), sends text to GPT-4o 
    for comparison, and returns a structured JSON result.
    """
    # 1. Input Validation and File Fetching
    if 'rfq_data' not in request.files or 'quote_data' not in request.files:
        return jsonify({'message': 'Missing one or both files in the request'}), 400
    rfq_file: FileStorage = request.files['rfq_data']
    quote_file: FileStorage = request.files['quote_data']
    if rfq_file.filename == '' or quote_file.filename == '':
        return jsonify({'message': 'One or more files have no filename'}), 400
    try:
        # --- CONSOLIDATED LOGIC: FILE PARSING (Inside the route) ---
        def parse_file_content_internal(file: FileStorage) -> str:
            """Reads and parses file content into a single string for AI analysis."""
            ext = get_file_extension(file.filename)
            file_bytes = file.stream.read()
            if ext in ['xlsx', 'xls', 'csv']:
                # Handle Excel/CSV using Pandas
                file_stream = io.BytesIO(file_bytes)
                if ext == 'csv':
                    df = pd.read_csv(file_stream)
                    return df.to_csv(index=False)
                else:
                    df_dict = pd.read_excel(file_stream, sheet_name=None)
                    all_sheets_csv = ""
                    for sheet_name, df_sheet in df_dict.items():
                        all_sheets_csv += f"\n--- Sheet: {sheet_name} ---\n" + df_sheet.to_csv(index=False)
                    return all_sheets_csv
            elif ext == 'pdf':
                # --- PDF Text Extraction Implementation ---
                reader = PdfReader(io.BytesIO(file_bytes))
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n" # Separate pages with double newline
                if not text.strip():
                    # Fallback for scanned/image-only PDFs
                    return f"PDF CONTENT ERROR: Text extraction failed for {file.filename} (might be scanned image). AI will only process the filename."
                return text.strip()
            elif ext in ['txt']:
                # Simple text file
                return file_bytes.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: .{ext}")
        # Execute parsing
        rfq_text = parse_file_content_internal(rfq_file)
        quote_text = parse_file_content_internal(quote_file)
        # --- CONSOLIDATED LOGIC: AI COMPARISON ---
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return jsonify({'message': "Configuration Error: OPENAI_API_KEY not set."}), 500
        client = OpenAI(api_key=api_key)
        json_structure = {
            "items_requested": "[list of items requested in RFQ]",
            "items_quoted": "[list of items quoted]",
            "differences_in_items": "[list of differences in items, including alternates or EOL items]",
            "discrepancies": "[list of pricing/term discrepancies]",
            "potential_issues": "[list of potential issues, e.g., EOL parts]",
            "summary": "summary of differences"
        }
        prompt_content = f"""
        You are an expert in analyzing RFQ and Quote documents.
        Given the following RFQ and Quote texts, identify discrepancies. 
        - Make sure the quoted items are the same as the requested items in the RFQ.
        - If items are different, highlight the differences.
        - If the items are alternatives, mention that.
        - If end-of-life (EOL) items are quoted, mention that.
        - If the RFQ contained an EOL item and an alternative was quoted, mention that.
        Provide the results **strictly** in the following JSON format: {json.dumps(json_structure, indent=2)}
        RFQ Text:
        ---
        {rfq_text}
        ---
        Quote Text:
        ---
        {quote_text}
        ---
        """
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", 
                 "content": "You are a specialized JSON output assistant. Your only task is to analyze the provided texts and return the analysis strictly as a single JSON object. DO NOT include any text outside the JSON object."},
                {"role": "user", "content": prompt_content}
            ],
            temperature=0,
            max_tokens=1500,
            response_format={"type": "json_object"}
        )
        json_string = response.choices[0].message.content
        comparison_result_dict = json.loads(json_string)
        # 4. Return result
        if "error" in comparison_result_dict:
             return jsonify({
                'message': 'AI Comparison Failed',
                'comparison_result': comparison_result_dict
            }), 500
        return jsonify({
            'message': f'Files processed and compared successfully: {rfq_file.filename} vs {quote_file.filename}',
            'comparison_result': comparison_result_dict
        }), 200
    except ValueError as ve:
        return jsonify({'message': f'File Processing Error (Unsupported Type/Content): {str(ve)}'}), 400
    except json.JSONDecodeError as e:
        return jsonify({'message': f'AI Model Error: Failed to parse JSON response. Details: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'message': f'OpenAI API Error: {str(e)}'}), 500
    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({'message': f'Internal Server Error: {str(e)}'}), 500
# ==============================================================
@app.route('/test_metadata', methods=['POST'])
def metadata_test():
    # If the other code sends JSON (e.g., requests.post(url, json=data))
    if request.is_json:
        data = request.get_json()
    # If the other code sends Form data (e.g., requests.post(url, data=data))
    else:
        data = request.form.to_dict()
    print(f"Received data: {data}")
    return jsonify({"status": "success", "received": data}), 200
@app.route('/assist')
def assist():
    user=session.get("user")
    if not user:
        return redirect(url_for('login'))
    user_name = user.get("displayName").replace(" ", "")
    tasks = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
    df = items_to_dataframe(tasks)
    # Filter for the current user and set fallbacks for Title
    user_items = []
    for t in tasks:
        if t.get("AssignedTo", "").replace(" ", "") == user_name:
            if not t.get('Title'):
                t['Title'] = t.get('ProjectName') or t.get('ProposalName') or t.get('Name') or t.get('ItemName') or 'Unnamed Task'
            user_items.append(t)
    print(f"User {user_name} has {len(user_items)} assigned items.", flush=True)
    return render_template("assist.html", user=user, tasks=user_items)
# @app.route("/line_items")
# def line_items():
#     if "user" not in session:
#         return redirect(url_for("login"))
#     user = session.get("user")
#     return render_template("line_items.html", user=user)
# ==============================================================
# PROCUREMENT AGENT API ROUTES
# ==============================================================
@app.route('/api/procurement/tasks', methods=['GET'])
def get_procurement_tasks():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    display_name = user.get("displayName", "").replace(" ", "")
    try:
        # Using the main Proposals list
        tasks_list = fetch_sharepoint_list(SITE_DOMAIN, SITE_PATH, LIST_NAME)
        email = user.get("mail") or user.get("userPrincipalName", "")
        # Filter comprehensively
        if is_admin(email):
            user_tasks = tasks_list
        else:
            user_tasks = [t for t in tasks_list if str(t.get("AssignedTo", "")).replace(" ", "").lower() == display_name.lower() or display_name.lower() in str(t.get("AssignedTo", "")).lower() or str(t.get("AssignedTo", "")).lower() in display_name.lower()]
        # Add basic priority and step data for the UI
        for t in user_tasks:
            if 'Priority' not in t: t['Priority'] = 'Medium'
            if 'steps_completed' not in t: t['steps_completed'] = 2
            if 'total_steps' not in t: t['total_steps'] = 5
            if 'DueAt' not in t: t['DueAt'] = t.get('BCD', 'N/A')
            # Create a robust Title fallback
            title_val = str(t.get('Title') or t.get('ProjectName') or t.get('ProposalName') or t.get('Name') or t.get('ItemName') or '').strip()
            if not title_val:
                title_val = "Unnamed Task"
            t['Title'] = title_val
        # Sort tasks by Created date descending (newest on top)
        def get_task_date(t):
            return t.get('Created', '')
        user_tasks.sort(key=get_task_date, reverse=True)
        return jsonify({"tasks": user_tasks})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/procurement/tasks/<task_id>', methods=['GET'])
def get_procurement_task_details(task_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    try:
        task_data = fetch_sharepoint_item_by_id(SITE_DOMAIN, SITE_PATH, LIST_NAME, task_id)
        if not task_data:
            return jsonify({"error": "Task not found"}), 404
        attachments_list = get_item_attachments(SITE_DOMAIN, SITE_PATH, LIST_NAME, task_id)
        return jsonify({
            "task": task_data,
            "attachments": attachments_list
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/procurement/analyze/<task_id>', methods=['GET', 'POST'])
def analyze_procurement_requirements(task_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    try:
        # Determine if we have a file in the request (manual upload)
        file_to_analyze = request.files.get('file')
        content_to_analyze = ""
        if file_to_analyze:
            print(f"[ANALYZE] Analyzing manually uploaded file: {file_to_analyze.filename}")
            ext = file_to_analyze.filename.split('.')[-1].lower()
            file_bytes = file_to_analyze.read()
            # Reuse parsing logic...
            content_to_analyze = _parse_file_content(ext, file_bytes, file_to_analyze.filename)
        else:
            # Traditional SharePoint path
            attachments = get_item_attachments(SITE_DOMAIN, SITE_PATH, LIST_NAME, task_id)
            if not attachments:
                return jsonify({"success": False, "error": "No attachments found to analyze."}), 400
            access_token = get_access_token()
            headers = {"Authorization": f"Bearer {access_token}"}
            for att in attachments:
                name = att['name'].lower()
                if any(name.endswith(ext) for ext in ['.pdf', '.xlsx', '.xls', '.docx', '.txt', '.csv']):
                    resp = requests.get(att['url'], headers=headers)
                    if resp.status_code == 200:
                        content_to_analyze += _parse_file_content(name.split('.')[-1], resp.content, name)
        if not content_to_analyze.strip():
            return jsonify({"success": False, "error": "Could not extract text from files."}), 400
        items = _run_requirement_analyzer(content_to_analyze)
        return jsonify({"success": True, "items": items})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500
def _parse_file_content(ext, file_bytes, filename):
    text = ""
    try:
        if ext == 'pdf':
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif ext in ['xlsx', 'xls']:
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            for sheet_name, df_sheet in df_dict.items():
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += df_sheet.to_csv(index=False) + "\n"
        elif ext == 'csv':
            df_att = pd.read_csv(io.BytesIO(file_bytes))
            text += df_att.to_csv(index=False) + "\n"
        elif ext == 'docx':
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text += "\n".join([p.text for p in doc.paragraphs]) + "\n"
        elif ext == 'txt':
            text += file_bytes.decode('utf-8', errors='ignore') + "\n"
    except Exception as e:
        print(f"Error parsing {filename}: {e}")
    return text
def _run_requirement_analyzer(content):
    try:
        from openai import OpenAI
        client = OpenAI(api_key=openai.api_key or os.getenv("OPENAI_API_KEY"))
        prompt = f"""Extract all items, products, or services needed for procurement from the following requirements description. 
Return them STRICTLY as a JSON list of objects, each with 'name' and 'type' (e.g., hardware, software, services).
Requirements:
{content[:15000]}
"""
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "You are a procurement expert. Return strictly a JSON list."},
                      {"role": "user", "content": prompt}]
        )
        res_text = response.choices[0].message.content
        match = re.search(r'\[.*\]', res_text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        return []
    except Exception as e:
        print(f"AI Analysis Error: {e}")
        return []
@app.route('/api/shared_projects/analyze/<project_id>', methods=['POST'])
def analyze_shared_project(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    try:
        project = get_shared_project_details(project_id)
        if not project: return jsonify({"error": "Project not found"}), 404
        content = ""
        # 1. From original task details
        if project.get('task_details'):
            task_id = project['task_details'].get('id')
            if task_id:
                try:
                    attachments = get_item_attachments(SITE_DOMAIN, SITE_PATH, LIST_NAME, task_id)
                    access_token = get_access_token()
                    headers = {"Authorization": f"Bearer {access_token}"}
                    for att in attachments:
                        name = att['name'].lower()
                        if any(name.endswith(ext) for ext in ['.pdf', '.xlsx', '.xls', '.docx', '.txt', '.csv']):
                            resp = requests.get(att['url'], headers=headers)
                            if resp.status_code == 200:
                                content += _parse_file_content(name.split('.')[-1], resp.content, name)
                except: pass
        # 2. From manually uploaded file in request
        file_to_analyze = request.files.get('file')
        if file_to_analyze:
            ext = file_to_analyze.filename.split('.')[-1].lower()
            content += _parse_file_content(ext, file_to_analyze.read(), file_to_analyze.filename)
        if not content.strip():
            return jsonify({"success": False, "error": "No content found to analyze for this shared project. Please upload a file."}), 400
        items = _run_requirement_analyzer(content)
        # Log this action in the project
        save_shared_session_message(project_id, "assistant", f"I've analyzed the project requirements and identified {len(items)} items for procurement.", "AI System")
        update_project_heartbeat(project_id, user_email) 
        return jsonify({"success": True, "items": items})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/procurement/find_distributors', methods=['POST'])
def procurement_find_distributors():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    items_to_search = data.get('items', [])
    try:
        categories = {}
        for item in items_to_search:
            cat = item.get('category') or item.get('Category') or item.get('type') or item.get('Type') or 'General'
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(item)
        all_distributors = []
        for cat, items_in_cat in categories.items():
            # Interally check the first item's name as representation of the category for past vendors
            rep_name = items_in_cat[0].get('name', cat)
            internal_results = search_item_distributors(rep_name)
            for res in internal_results:
                history = res.get('purchase_history', [])
                for h in history:
                    vendor = h.get('Vendor') or h.get('Distributor')
                    if vendor:
                        all_distributors.append({
                            "name": vendor,
                            "email": h.get("Email", "Not in DB"),
                            "source": "internal",
                            "item": cat,
                            "items_list": items_in_cat
                        })
            web_data = search_web(f"authorized distributors for {cat} in UAE")
            try:
                import json as py_json
                web_list = py_json.loads(web_data)
                if isinstance(web_list, list):
                    for w in web_list[:2]:
                        all_distributors.append({
                            "name": w.get('title', 'Unknown Distributor'),
                            "email": w.get('email', 'Not Found'), 
                            "source": "web",
                            "item": cat,
                            "items_list": items_in_cat,
                            "link": w.get('href', '#')
                        })
            except:
                pass
        unique_list = []
        seen_names = set()
        for d in all_distributors:
            if d['name'].lower() not in seen_names:
                unique_list.append(d)
                seen_names.add(d['name'].lower())
        return jsonify({"distributors": unique_list})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/procurement/draft_email', methods=['POST'])
def procurement_draft_email():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    distributor_info = data.get('distributor')
    if not distributor_info:
        return jsonify({"error": "Missing distributor info"}), 400
    try:
        dist_name = distributor_info.get('name')
        items = data.get('items', [])
        items_req = "my task"
        if items:
             items_text = "\n".join([f"- {i.get('name') or i.get('Name', 'Item')} (Qty: {i.get('quantity') or i.get('Quantity', '1')} {i.get('unit') or i.get('Unit', '')})" for i in items])
             items_req = f"the following items:\n{items_text}"
        custom_prompt = data.get('customPrompt')
        current_draft = data.get('currentDraft')
        if custom_prompt and current_draft:
            prompt_email = f"You are a procurement assistant. Rewrite the following email draft. \n\nCURRENT DRAFT:\n{current_draft}\n\nUSER INSTRUCTIONS:\n{custom_prompt}\n\nReturn as JSON with 'subject' and 'body'. The body should be formatted in clean HTML (e.g., using <br> or <p> tags)."
        else:
            prompt_email = f"Draft a professional procurement inquiry email to {dist_name}. Inquire about availability, lead times, and pricing for {items_req}. Return as JSON with 'subject' and 'body'. The body must consist of professional HTML."
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        draft_res = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a procurement assistant drafting emails. Return JSON only."},
                {"role": "user", "content": prompt_email}
            ],
            temperature=0.7
        )
        res_text = draft_res.choices[0].message.content
        try:
            import json as py_json
            match_json = re.search(r'\{.*\}', res_text, re.DOTALL)
            draft_data = py_json.loads(match_json.group(0)) if match_json else py_json.loads(res_text)
            # Normalize keys to lowercase for robust frontend usage
            draft_data = {k.lower(): v for k, v in draft_data.items()}
            draft_data['to'] = distributor_info.get('email') if distributor_info.get('email') not in ["Searching...", "Not Found", None, ""] else ""
        except:
            draft_data = {
                "subject": f"Inquiry for items - {dist_name}",
                "body": f"Dear {dist_name},\n\nWe are interested in procuring {items_req}. Please provide a quote.",
                "to": distributor_info.get('email') if distributor_info.get('email') not in ["Searching...", "Not Found", None, ""] else ""
            }
        return jsonify({"draft": draft_data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/procurement/feedback', methods=['POST'])
def procurement_feedback():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    items = data.get('items', [])
    distributors = data.get('distributors', [])
    is_true_data = data.get('is_true_data', False)
    notes = data.get('notes', "")
    user = session["user"]
    user_email = user.get("mail") or user.get("userPrincipalName")
    from cosmos import save_procurement_feedback
    success = save_procurement_feedback(user_email, items, distributors, is_true_data, notes)
    if success:
        return jsonify({"success": True, "message": "Feedback saved successfully."})
    else:
        return jsonify({"success": False, "error": "Failed to save feedback."}), 500
# ==============================================================
# MAIL DASHBOARD ROUTES
# ==============================================================
@app.route('/mails')
def mail_dashboard():
    if "user" not in session:
        return redirect(url_for('login'))
    user = session["user"]
    return render_template('mail_dashboard.html', user=user)
@app.route('/api/mails')
def get_mails():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    # Use application token instead of delegated token
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    folder = request.args.get('folder', 'inbox')
    well_known_folders = ['inbox', 'sentitems', 'drafts', 'deleteditems', 'junkemail', 'archive']
    try:
        if folder in well_known_folders:
            url = f"{GRAPH_API_ENDPOINT}/users/{email}/mailFolders/{folder}/messages?$top=50&$select=id,subject,bodyPreview,sender,receivedDateTime,isRead,inferenceClassification"
        else:
            url = f"{GRAPH_API_ENDPOINT}/users/{email}/messages?$top=50&$select=id,subject,bodyPreview,sender,receivedDateTime,isRead,inferenceClassification"
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            return jsonify(response.json())
        else:
            return jsonify({"error": "Failed to fetch emails via App Token", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/mails/<message_id>')
def get_mail_details(message_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    try:
        url = f"{GRAPH_API_ENDPOINT}/users/{email}/messages/{message_id}?$select=id,subject,body,sender,toRecipients,receivedDateTime,isRead"
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            return jsonify(response.json())
        else:
            return jsonify({"error": "Failed to fetch email details via App Token", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/mails/<message_id>', methods=['DELETE'])
def delete_mail(message_id):
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {"Authorization": f"Bearer {access_token}"}
    try:
        url = f"{GRAPH_API_ENDPOINT}/users/{email}/messages/{message_id}"
        response = requests.delete(url, headers=headers)
        if response.status_code == 204:
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Failed to delete email", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/mails/<message_id>/<action>', methods=['POST'])
def mail_action(message_id, action):
    # action can be: reply, replyAll, forward
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    if action not in ['reply', 'replyAll', 'forward']:
        return jsonify({"error": "Invalid action"}), 400
    user = session["user"]
    email_addr = user.get("mail") or user.get("userPrincipalName")
    data = request.get_json()
    comment = data.get('comment', '')
    to_recipients = data.get('toRecipients', []) # only needed strictly for forward or adding people
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    payload = {
        "message": {
            "body": {
                "contentType": "HTML",
                "content": comment
            }
        }
    }
    if action == 'forward' and to_recipients:
        payload["message"]["toRecipients"] = [{"emailAddress": {"address": addr}} for addr in to_recipients]
    try:
        url = f"{GRAPH_API_ENDPOINT}/users/{email_addr}/messages/{message_id}/{action}"
        response = requests.post(url, headers=headers, json=payload)
        if response.status_code == 202:
            return jsonify({"success": True})
        else:
            return jsonify({"error": f"Failed to {action} email", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/mails/send', methods=['POST'])
def send_mail():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    user_email = user.get("mail") or user.get("userPrincipalName")
    if request.is_json:
        data = request.get_json()
        to_email = data.get('to')
        subject = data.get('subject')
        body_content = data.get('body')
        task_id = data.get('task_id')
        session_id = data.get('session_id')
        files = []
    else:
        to_email = request.form.get('to')
        subject = request.form.get('subject')
        body_content = request.form.get('body')
        task_id = request.form.get('task_id')
        session_id = request.form.get('session_id')
        files = request.files.getlist('attachments')
    if not all([to_email, subject, body_content]):
        return jsonify({"error": "Missing required fields: to, subject, or body"}), 400
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    tracking_id = str(uuid4())
    parsed_body = body_content
    if task_id:
        parsed_body += f"<br><br><span style='color: transparent; font-size: 0px;'>REF:{tracking_id}</span>"
    mail_payload = {
        "message": {
            "subject": subject,
            "body": {
                "contentType": "HTML",
                "content": parsed_body
            },
            "toRecipients": [
                {
                    "emailAddress": {
                        "address": to_email.strip()
                    }
                }
            ]
        },
        "saveToSentItems": "true"
    }
    if files:
        import base64
        attachments_list = []
        for f in files:
            if f.filename:
                file_bytes = f.read()
                b64_content = base64.b64encode(file_bytes).decode('utf-8')
                attachments_list.append({
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": f.filename,
                    "contentBytes": b64_content
                })
        if attachments_list:
            mail_payload["message"]["hasAttachments"] = True
            mail_payload["message"]["attachments"] = attachments_list
    try:
        url = f"{GRAPH_API_ENDPOINT}/users/{user_email}/sendMail"
        response = requests.post(url, headers=headers, json=mail_payload)
        if response.status_code == 202:
            if task_id:
                save_tracked_email(task_id, session_id, to_email, subject, tracking_id, user_email, body_content)
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Failed to send email via App Token", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/api/mails/draft/save', methods=['POST'])
def save_mail_draft():
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    user = session["user"]
    user_email = user.get("mail") or user.get("userPrincipalName")
    if request.is_json:
        data = request.get_json()
        to_email = data.get('to')
        subject = data.get('subject')
        body_content = data.get('body')
        draft_id = data.get('draftId')
        files = []
    else:
        to_email = request.form.get('to')
        subject = request.form.get('subject')
        body_content = request.form.get('body')
        draft_id = request.form.get('draftId')
        files = request.files.getlist('attachments')
    if not all([to_email, subject, body_content]):
        return jsonify({"error": "Missing required fields: to, subject, or body"}), 400
    try:
        from sharepoint_items import get_access_token
        access_token = get_access_token()
    except Exception as e:
        return jsonify({"error": "Failed to acquire app token", "details": str(e)}), 500
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    mail_payload = {
        "subject": subject,
        "body": {
            "contentType": "HTML",
            "content": body_content
        },
        "toRecipients": [
            {
                "emailAddress": {
                    "address": to_email.strip()
                }
            }
        ]
    }
    if files:
        import base64
        attachments_list = []
        for f in files:
            if f.filename:
                file_bytes = f.read()
                b64_content = base64.b64encode(file_bytes).decode('utf-8')
                attachments_list.append({
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": f.filename,
                    "contentBytes": b64_content
                })
        if attachments_list:
            mail_payload["hasAttachments"] = True
            mail_payload["attachments"] = attachments_list
    try:
        if draft_id and draft_id != "null" and draft_id != "undefined":
            url = f"{GRAPH_API_ENDPOINT}/users/{user_email}/messages/{draft_id}"
            response = requests.patch(url, headers=headers, json=mail_payload)
            if response.status_code == 200:
                data = response.json()
                return jsonify({"success": True, "webLink": data.get("webLink", ""), "id": data.get("id")})
            else:
                return jsonify({"error": "Failed to update draft", "details": response.text}), response.status_code
        else:
            url = f"{GRAPH_API_ENDPOINT}/users/{user_email}/messages"
            response = requests.post(url, headers=headers, json=mail_payload)
            if response.status_code == 201:
                data = response.json()
                return jsonify({"success": True, "webLink": data.get("webLink", ""), "id": data.get("id")})
            else:
                return jsonify({"error": "Failed to save draft", "details": response.text}), response.status_code
    except Exception as e:
        return jsonify({"error": str(e)}), 500
# ==============================================================
# COLLABORATION & SHARED PROJECTS API
# ==============================================================
@app.route('/api/users/search', methods=['GET'])
def search_users():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    query = request.args.get('q', '').lower()
    if len(query) < 2: return jsonify({"users": []})
    try:
        # Search via Microsoft Graph for team members
        from sharepoint_items import get_access_token
        token = get_access_token()
        headers = {"Authorization": f"Bearer {token}"}
        url = f"{GRAPH_API_ENDPOINT}/users?$filter=startswith(displayName,'{query}') or startswith(mail,'{query}') or startswith(userPrincipalName,'{query}')&$select=displayName,mail,userPrincipalName&$top=10"
        resp = requests.get(url, headers=headers)
        if resp.status_code == 200:
            users_data = resp.json().get('value', [])
            return jsonify({"users": users_data})
        return jsonify({"users": []})
    except Exception as e:
        print(f"[API ERROR] search_users: {e}")
        return jsonify({"users": []})
@app.route('/api/shared_projects', methods=['GET', 'POST'])
def handle_shared_projects():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    if request.method == 'GET':
        projects = get_shared_projects_for_user(user_email)
        return jsonify({"projects": projects})
    elif request.method == 'POST':
        data = request.get_json()
        task_data = data.get('task')
        if not task_data: return jsonify({"error": "Missing task data"}), 400
        project_id = create_shared_project(task_data, user_email)
        if project_id:
            return jsonify({"success": True, "project_id": project_id})
        return jsonify({"error": "Failed to create project"}), 500
@app.route('/api/shared_projects/<project_id>', methods=['GET'])
def get_project(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    project = get_shared_project_details(project_id)
    if not project: return jsonify({"error": "Project not found"}), 404
    # Check if user is a collaborator
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    if user_email not in project.get("collaborators", []):
        return jsonify({"error": "Access denied"}), 403
    activity = get_shared_project_activity(project_id)
    return jsonify({"project": project, "activity": activity})
@app.route('/api/shared_projects/<project_id>/invite', methods=['POST'])
def invite_to_project(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    data = request.get_json()
    invitee_emails = data.get('emails', [])
    project = get_shared_project_details(project_id)
    if not project: return jsonify({"error": "Project not found"}), 404
    success_count = 0
    for email in invitee_emails:
        if invite_collaborator(project_id, user_email, email):
            # Send invitation email using Graph API
            try:
                from sharepoint_items import get_access_token
                token = get_access_token()
                headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
                mail_payload = {
                    "message": {
                        "subject": f"Collaboration Invite: {project['task_details'].get('Title')}",
                        "body": {
                            "contentType": "HTML",
                            "content": f"""
                                <h2>Collaboration Invite</h2>
                                <p><strong>{user_email}</strong> has invited you to collaborate on the task: 
                                <strong>{project['task_details'].get('Title')}</strong>.</p>
                                <p>Please open the Hamdaz Chatbot to accept the invitation.</p>
                                <hr/>
                                <p><a href="{request.host_url}personal_assistant">Open Dashboard</a></p>
                            """
                        },
                        "toRecipients": [{"emailAddress": {"address": email}}]
                    }
                }
                requests.post(f"{GRAPH_API_ENDPOINT}/users/{user_email}/sendMail", headers=headers, json=mail_payload)
            except Exception as e:
                print(f"[MAIL ERROR] Failed to send invite email to {email}: {e}")
            success_count += 1
    return jsonify({"success": True, "invited": success_count})
@app.route('/api/shared_projects/<project_id>/accept', methods=['POST'])
def accept_invite(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    data = request.get_json()
    notification_id = data.get('notification_id')
    if accept_collaboration_invite(notification_id, user_email):
        return jsonify({"success": True})
    return jsonify({"error": "Failed to accept invite"}), 500
@app.route('/api/shared_projects/<project_id>/chat', methods=['POST'])
def shared_chat(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    # Support both JSON and Multipart Form Data
    if request.is_json:
        data = request.get_json()
        user_prompt = data.get('message')
        files = []
        is_analysis = str(data.get('is_analysis', 'false')).lower() == 'true'
    else:
        user_prompt = request.form.get('message')
        files = request.files.getlist('files')
        is_analysis = request.form.get('is_analysis', 'false').lower() == 'true'
    files_text = ""
    if files:
        for file in files:
            if file.filename == "": continue
            try:
                ext = file.filename.split('.')[-1].lower()
                file_bytes = file.read()
                if ext == 'pdf':
                    from pypdf import PdfReader
                    reader = PdfReader(io.BytesIO(file_bytes))
                    for page in reader.pages:
                        files_text += (page.extract_text() or "") + "\n"
                elif ext in ['xlsx', 'xls']:
                    df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                    for sheet_name, df_sheet in df_dict.items():
                        files_text += f"\n--- Sheet: {sheet_name} ---\n"
                        files_text += df_sheet.to_csv(index=False) + "\n"
                elif ext == 'csv':
                    df_att = pd.read_csv(io.BytesIO(file_bytes))
                    files_text += df_att.to_csv(index=False) + "\n"
                elif ext == 'docx':
                    import docx
                    doc = docx.Document(io.BytesIO(file_bytes))
                    files_text += "\n".join([p.text for p in doc.paragraphs]) + "\n"
                elif ext == 'txt':
                    files_text += file_bytes.decode('utf-8', errors='ignore') + "\n"
            except Exception as e:
                print(f"[SHARED CHAT] Error parsing file {file.filename}: {e}")
                files_text += f"\n[Error parsing file {file.filename}: {str(e)}]\n"
    project = get_shared_project_details(project_id)
    if not project: return jsonify({"error": "Project not found"}), 404
    # Save user message (Prompt only)
    save_shared_session_message(project_id, "user", user_prompt, user_email)
    # Inject collaboration context into the AI
    other_collaborators = [c for c in project.get("collaborators", []) if c != user_email]
    context = f"\n\nYou are in a COLLABORATIVE SESSION for the project: '{project['task_details'].get('Title')}'.\n"
    context += f"Collaborators: {', '.join(project.get('collaborators', []))}.\n"
    if other_collaborators:
        context += f"Inform the user about what others might be doing if relevant. Active users: {', '.join(get_project_presence(project_id))}.\n"
    # Run assistant
    history = project.get("messages", [])
    clean_history = [{"role": m["role"], "content": m["content"]} for m in history]
    # Combined prompt with context
    full_prompt = context + (user_prompt or "")
    system_instr = ""
    if is_analysis:
        system_instr = "User has uploaded a procurement document. ANALYZE it and return the required items STRICTLY as a JSON list of objects with 'name' and 'type' properties. DO NOT add any conversational text."
    ai_response = run_personal_assistant(user_email, full_prompt, files_text=files_text, chat_history=clean_history, system_instr=system_instr)
    # Save AI response
    save_shared_session_message(project_id, "assistant", ai_response, "AI")
    return jsonify({"response": ai_response})
@app.route('/api/shared_projects/<project_id>/heartbeat', methods=['POST'])
def project_heartbeat(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    update_project_heartbeat(project_id, user_email)
    return jsonify({"success": True})
@app.route('/api/shared_projects/<project_id>/presence', methods=['GET'])
def project_presence(project_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    active_users = get_project_presence(project_id)
    return jsonify({"active_users": active_users})
@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    user_email = (session["user"].get("mail") or session["user"].get("userPrincipalName", "")).lower()
    notifs = get_user_notifications(user_email)
    return jsonify({"notifications": notifs})
@app.route('/api/notifications/<notification_id>/read', methods=['POST'])
def read_notification(notification_id):
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    mark_notification_read(notification_id)
    return jsonify({"success": True})
@app.route('/api/quotes/status', methods=['POST'])
def api_quote_status():
    if 'user' not in session: return jsonify({'success': False}), 401
    data = request.json
    tracking_id = data.get('tracking_id')
    task_id = data.get('task_id')
    status = data.get('status')
    from cosmos import task_supplier_quotes_container, update_task_supplier_quote_status, save_task_supplier_quote, tracked_emails_container
    if not task_supplier_quotes_container: return jsonify({'success': False}), 500
    query = "SELECT * FROM c WHERE c.tracking_id = @trackId AND c.task_id = @taskId"
    res = list(task_supplier_quotes_container.query_items(query=query, parameters=[{'name':'@trackId','value':tracking_id},{'name':'@taskId','value':task_id}], enable_cross_partition_query=True))
    if not res:
        # Fallback: Check if it exists in tracked_emails and has AI data
        print(f"[STATUS] Quote {tracking_id} not in dedicated container. Checking tracked_emails fallback...")
        tracked_res = list(tracked_emails_container.query_items(query="SELECT * FROM c WHERE c.id = @id", parameters=[{'name':'@id','value':tracking_id}], enable_cross_partition_query=True))
        if tracked_res:
            t_doc = tracked_res[0]
            ai_data = t_doc.get('ai_parsed_data')
            if ai_data:
                # Synthesize the record in task_supplier_quotes
                save_task_supplier_quote(task_id, tracking_id, t_doc.get('to_email'), t_doc.get('summary'), ai_data)
                # Re-query
                res = list(task_supplier_quotes_container.query_items(query=query, parameters=[{'name':'@trackId','value':tracking_id},{'name':'@taskId','value':task_id}], enable_cross_partition_query=True))
            else:
                return jsonify({'success': False, 'message': 'No AI data found for this email to shortlist.'}), 400
        else:
            return jsonify({'success': False, 'message': 'Quote not found'}), 404
    if not res: return jsonify({'success': False, 'message': 'Failed to synthesize quote record'}), 500
    quote_id = res[0]['id']
    success = update_task_supplier_quote_status(quote_id, task_id, status)
    return jsonify({'success': success})
@app.route('/api/extracted_quotes/<task_id>', methods=['GET'])
def api_get_extracted_quotes(task_id):
    if 'user' not in session: return jsonify({'success': False}), 401
    from cosmos import get_task_supplier_quotes
    # Only pull quotes that the user actively chose to add to the shortlist collection
    quotes_docs = get_task_supplier_quotes(task_id, status_filter='collected')
    quotes = []
    for q in quotes_docs:
        quotes.append({
            'quote_id': q.get('id'),
            'tracking_id': q.get('tracking_id'),
            'from_email': q.get('supplier_email'),
            'items': q.get('items', []),
            'summary': q.get('summary'),
            'created_at': q.get('created_at')
        })
    return jsonify({'success': True, 'data': quotes})
@app.route('/api/quotes/download/<tracking_id>', methods=['GET'])
def api_download_quote(tracking_id):
    if 'user' not in session: return jsonify({'error': 'Unauthorized'}), 401
    from cosmos import tracked_emails_container
    if not tracked_emails_container: return jsonify({'error': 'DB Error'}), 500
    # Find the document by cross partition query (since partition key is task_id)
    res = list(tracked_emails_container.query_items(query='SELECT * FROM c WHERE c.id = @id', parameters=[{'name':'@id', 'value':tracking_id}], enable_cross_partition_query=True))
    if not res: return jsonify({'error': 'Not found'}), 404
    doc = res[0]
    parsed = doc.get('ai_parsed_data')
    if not parsed: return jsonify({'error': 'No quote data'}), 404
    from quote_generator import generate_commercial_proposal_docx
    doc_io = generate_commercial_proposal_docx(parsed, tracking_id)
    from flask import send_file
    return send_file(doc_io, download_name=f'Commercial_Proposal_{tracking_id[:8].upper()}.docx', as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
@app.route('/api/chat/init', methods=['POST'])
def api_init_chat():
    if 'user' not in session: return jsonify({'error': 'Unauthorized'}), 401
    data = request.json or {}
    session_id = data.get('session_id')
    task_id = data.get('task_id')
    task_title = data.get('title', 'Procurement Context')
    if not session_id: return jsonify({'error': 'Missing session_id'}), 400
    user_email = session['user'].get('mail') or session['user'].get('userPrincipalName')
    # Use internal cosmos func to init session without real user message
    from cosmos import save_session_message
    # We save a hidden system initialized flag
    save_session_message(session_id, user_email, 'system', f'Session initialized for {task_title}', title=task_title, agent_type='procurement', task_id=task_id)
    return jsonify({'success': True})
# ==============================================================
# LEAVE MANAGEMENT ROUTES
# ==============================================================
@app.route("/leave")
def leave_page():
    if "user" not in session:
        return redirect(url_for("login"))
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    username = user.get("displayName", "").replace(" ", "")
    return render_template("pages/leave.html", user=user, email=email, username=username)
@app.route("/api/leave/ongoing_proposals", methods=["GET"])
def api_leave_ongoing_proposals():
    """Returns the current user s active/ongoing proposals for handoff preview."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user = session["user"]
    username = user.get("displayName", "").replace(" ", "")
    try:
        proposals = get_ongoing_proposals_for_user(username)
        return jsonify({"success": True, "proposals": proposals})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/priority_user", methods=["GET"])
def api_leave_priority_user():
    """Returns the Priority-1 user from useranalytics for auto-assign preview."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user = session["user"]
    username = user.get("displayName", "").replace(" ", "")
    try:
        priority_user = get_priority_one_user(exclude_username=username)
        return jsonify({"success": True, "priority_user": priority_user})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/users", methods=["GET"])
def api_leave_users():
    """Returns list of all org users for the manual handoff dropdown."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    try:
        token = get_access_token()
        users_raw = get_all_users(token)
        # users_raw is a dict {id: displayName}
        users_list = [{"id": k, "displayName": v} for k, v in users_raw.items() if v]
        return jsonify({"success": True, "users": users_list})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/submit", methods=["POST"])
def api_leave_submit():
    """
    Handles leave request submission:
    1. Optionally adds user to SP excludeusers list
    2. Optionally handoffs proposals (auto or manual)
    3. Saves leave record to Cosmos DB
    """
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    username = user.get("displayName", "").replace(" ", "")
    data = request.get_json() or {}
    leave_start = data.get("leave_start")
    leave_end = data.get("leave_end")
    continue_assign = bool(data.get("continue_assign", False))
    handoff_enabled = bool(data.get("handoff_enabled", False))
    handoff_mode = data.get("handoff_mode", "auto")  # "auto" or "manual"
    manual_user = data.get("handoff_to_user", "")     # username for manual mode
    leave_type = data.get("leave_type", "full_day")    # "full_day", "first_half", "second_half"
    if not leave_start or not leave_end:
        return jsonify({"success": False, "error": "Leave start and end dates are required"}), 400
    # Start: Global Concurrency Limit validation
    leave_category = data.get("leave_category", "Casual")
    leave_reason = data.get("leave_reason", "")
    from cosmos import get_leave_settings, get_max_concurrent_leave_count, save_leave_request
    settings = get_leave_settings() or {}
    max_limit = settings.get("max_concurrent_limit", 0)
    auto_status = "active"
    reject_reason = ""
    current_peak = 0
    auto_approval_enabled = settings.get("auto_approval_enabled", True)
    waitlist_enabled = settings.get("waitlist_enabled", False)
    
    if auto_approval_enabled:
        if max_limit > 0:
            current_peak = get_max_concurrent_leave_count(leave_start, leave_end)
            if current_peak >= max_limit:
                auto_status = "WL" if waitlist_enabled else "rejected"
                reject_reason = f"Global leave capacity reached ({max_limit} active leaves) for the requested dates. Status changed to {auto_status}."
        print(f"[LEAVE SUBMIT] Concurrency Check: Peak={current_peak if max_limit > 0 else 'N/A'}, Limit={max_limit}. Status: {auto_status}")
    else:
        auto_status = "pending"
        print(f"[LEAVE SUBMIT] Auto-approval disabled. Status: {auto_status}")
    handoff_to_user = None
    proposals_transferred = []
    handoff_results = {}
    try:
        # Only do exclude/handoff if auto-approved
        if auto_status == "active":
            # ---- Step 1: Exclude user from task assignment ----
            if not continue_assign:
                add_user_to_excludelist(username)
                print(f"[LEAVE] Added {username} to excludeusers SP list.")
            # ---- Step 2: Handoff proposals ----
            if handoff_enabled:
                proposals = get_ongoing_proposals_for_user(username)
                proposal_ids = [p["id"] for p in proposals]
                proposals_transferred = [p["Title"] for p in proposals]
                if handoff_mode == "auto":
                    handoff_to_user = get_priority_one_user(exclude_username=username)
                else:
                    handoff_to_user = manual_user
                if handoff_to_user and proposal_ids:
                    handoff_results = handoff_proposals_to_user(username, handoff_to_user, proposal_ids)
                    print(f"[LEAVE] Handoff: {len(handoff_results['success'])} proposals moved to {handoff_to_user}")
                elif not proposal_ids:
                    print("[LEAVE] No proposals to handoff.")
                else:
                    print("[LEAVE] Could not determine handoff user.")
        else:
            print(f"[LEAVE] Leave handled manually or waitlisted ({auto_status}). Skipping exclude and handoff.")
        # ---- Step 3: Save to Cosmos DB ----
        doc_id = save_leave_request(
            user_email=email,
            username=username,
            leave_start=leave_start,
            leave_end=leave_end,
            continue_assign=continue_assign,
            handoff_enabled=handoff_enabled,
            handoff_mode=handoff_mode,
            handoff_to_user=handoff_to_user,
            proposals_transferred=proposals_transferred,
            leave_type=leave_type,
            status=auto_status,
            leave_category=leave_category,
            leave_reason=leave_reason,
            reviewed_by="System (Auto)" if auto_status in ["active", "rejected"] else None
        )
        if doc_id:
            # SEND NOTIFICATION EMAILS
            from cosmos import get_leave_settings
            settings = get_leave_settings() or {}
            hr_email = settings.get("hr_email")
            # 1. Notify HR
            hr_subject = f"Leave Request - {username} - {auto_status.upper()}"
            _status_str = "Waitlisted (WL)" if auto_status == "WL" else auto_status.upper()
            hr_body = f"""
                <p>User <b>{username}</b> ({email}) submitted a leave request.</p>
                <p><b>Dates:</b> {leave_start} to {leave_end}<br>
                <b>Type:</b> {leave_type.replace('_',' ')}<br>
                <b>Category:</b> {leave_category}<br>
                <b>Reason:</b> {leave_reason or 'No reason provided'}</p>
                <p><b>System Status:</b> <b>{_status_str}</b><br>{reject_reason}</p>
            """
            send_graph_email(hr_email, hr_subject, hr_body)
            # 2. Notify User if rejected or waitlisted
            if auto_status == "rejected":
                user_subject = "Leave Request Auto-Rejected"
                user_body = f"Hello {username},<br><br>Your leave request from {leave_start} to {leave_end} was automatically rejected by the system because the concurrent limit has been reached.<br><br>Please contact HR if you have questions."
                send_graph_email(email, user_subject, user_body)
            elif auto_status == "WL":
                user_subject = "Leave Request Waitlisted"
                user_body = f"Hello {username},<br><br>Your leave request from {leave_start} to {leave_end} has been placed on the <b>Waitlist</b> because the concurrent leave limit is currently full. Admin/HR will review your request manually.<br><br>Please contact HR if you have urgent questions."
                send_graph_email(email, user_subject, user_body)
            return jsonify({
                "success": True, 
                "doc_id": doc_id,
                "handoff_to": handoff_to_user,
                "proposals_transferred": proposals_transferred,
                "status": auto_status,
                "reject_reason": reject_reason
            })
        else:
            return jsonify({"success": False, "error": "Database error"}), 500
    except Exception as e:
        print(f"[ERROR] api_leave_submit: {e}", flush=True)
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/availability", methods=["GET"])
def api_leave_availability():
    """Returns total slots and currently free slots for today."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    try:
        from cosmos import get_leave_settings, get_max_concurrent_leave_count
        settings = get_leave_settings() or {}
        max_limit = settings.get("max_concurrent_limit", 3)
        today = datetime.now().strftime("%Y-%m-%d")
        # Reuse peak count logic for just today
        occupied = get_max_concurrent_leave_count(today, today)
        free = max(0, int(max_limit) - occupied)
        return jsonify({
            "success": True,
            "total_slots": max_limit,
            "occupied_slots": occupied,
            "free_slots": free
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/settings", methods=["GET"])
def api_leave_settings_public():
    """Allows regular users to fetch non-sensitive leave settings."""
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    from cosmos import get_leave_settings
    settings = get_leave_settings() or {}
    return jsonify({
        "success": True,
        "settings": {
            "max_concurrent_limit": settings.get("max_concurrent_limit", 3),
            "hr_email": settings.get("hr_email"),
            "auto_approval_enabled": settings.get("auto_approval_enabled", True),
            "waitlist_enabled": settings.get("waitlist_enabled", False)
        }
    })
@app.route("/api/leave/history", methods=["GET"])
def api_leave_history():
    """Returns leave history for the current user from Cosmos DB."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    if not email:
        return jsonify({"success": False, "error": "User email not found"}), 400
    try:
        from cosmos import get_leave_history_for_user
        history = get_leave_history_for_user(email)
        return jsonify({"success": True, "history": history})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/leave/cancel/<doc_id>", methods=["POST"])
def api_leave_cancel(doc_id):
    """
    Cancels a leave request:
    1. Marks it cancelled in Cosmos
    2. Removes user from SP excludeusers (so they resume receiving tasks)
    """
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    user = session["user"]
    email = user.get("mail") or user.get("userPrincipalName")
    username = user.get("displayName", "").replace(" ", "")
    try:
        from cosmos import cancel_leave_request
        cancel_leave_request(doc_id, email)
        # Always attempt to remove from excludeusers on cancel
        remove_user_from_excludelist(username)
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/admin/leaves")
def admin_leaves_page():
    """Admin Leave Dashboard - only accessible by admins."""
    if "user" not in session:
        return redirect(url_for("login"))
    email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(email):
        return redirect(url_for("leave_page"))
    return render_template("pages/admin_leaves.html", user=session["user"])
@app.route("/api/admin/leaves", methods=["GET"])
def api_admin_leaves():
    """Returns all leave records across all users. Admin-only."""
    if "user" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401
    email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(email):
        return jsonify({"success": False, "error": "Admin access required"}), 403
    try:
        from cosmos import get_all_leaves
        leaves = get_all_leaves()
        return jsonify({"success": True, "leaves": leaves})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
# ==============================================================
# ==============================================================
# ==============================================================
# LEAVE SYSTEM NOTIFICATIONS & ADMIN APIs
# ==============================================================
def send_graph_email(to_email, subject, html_body):
    """Sends a system email regarding leave requests using the MS Graph App Token."""
    from sharepoint_items import get_access_token
    try:
        token = get_access_token()
        if not token:
            print("[LEAVE EMAIL ERROR] No access token available.")
            return False
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        # We send from a system-designated admin/HR account for now.
        from cosmos import get_leave_settings
        _s = get_leave_settings() or {}
        from_email = _s.get("hr_email", "hr@hamdaz.com")
        message = {
            "message": {
                "subject": subject,
                "body": {
                    "contentType": "html",
                    "content": html_body
                },
                "toRecipients": [{"emailAddress": {"address": to_email}}]
            }
        }
        url = f"{GRAPH_API_ENDPOINT}/users/{from_email}/sendMail"
        response = requests.post(url, headers=headers, json=message)
        if response.status_code == 202:
            print(f"[LEAVE EMAIL] Sent '{subject}' to {to_email}")
            return True
        else:
            print(f"[LEAVE EMAIL ERROR] {response.status_code} - {response.text}")
            return False
    except Exception as e:
        print(f"[LEAVE EMAIL EXCEPTION] {e}")
        return False
@app.route("/api/admin/leave/settings", methods=["GET", "POST"])
def api_admin_leave_settings():
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(email): return jsonify({"success": False, "error": "Access Denied"}), 403
    from cosmos import get_leave_settings, save_leave_setting
    if request.method == "GET":
        settings = get_leave_settings()
        return jsonify({"success": True, "settings": settings or {}})
    if request.method == "POST":
        data = request.json
        setting_data = {
            "id": "leave_limit",
            "setting_type": "config",
            "max_concurrent_limit": int(data.get("max_concurrent_limit", 3)),
            "hr_email": data.get("hr_email"),
            "auto_approval_enabled": data.get("auto_approval_enabled", True),
            "waitlist_enabled": data.get("waitlist_enabled", False),
            "updated_by": email,
            "updated_at": datetime.utcnow().isoformat()
        }
        res = save_leave_setting(setting_data)
        return jsonify({"success": res is not None})
@app.route("/api/admin/leave/holidays", methods=["GET", "POST"])
def api_admin_leave_holidays():
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    from cosmos import get_holidays, save_holiday
    if request.method == "GET":
        holidays = get_holidays()
        return jsonify({"success": True, "holidays": holidays})
    if request.method == "POST":
        if not is_admin(email): return jsonify({"success": False, "error": "Access Denied"}), 403
        data = request.json
        doc_id = save_holiday(
            title=data.get("title"),
            date_str=data.get("date"),
            end_date_str=data.get("end_date") or data.get("date"),
            holiday_type=data.get("type", "holiday"),
            description=data.get("description", ""),
            created_by=email
        )
        return jsonify({"success": doc_id is not None, "id": doc_id})
@app.route("/api/admin/leave/holidays/<doc_id>", methods=["DELETE"])
def api_admin_delete_holiday(doc_id):
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(email): return jsonify({"success": False, "error": "Access Denied"}), 403
    from cosmos import delete_holiday
    res = delete_holiday(doc_id)
    return jsonify({"success": res})
@app.route("/api/admin/leave/approve/<doc_id>", methods=["POST"])
def api_admin_approve_leave(doc_id):
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    admin_email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(admin_email): return jsonify({"success": False, "error": "Access Denied"}), 403
    data = request.json or {}
    remarks = data.get("remarks", "Approved via Dashboard")
    target_user_email = data.get("user_email")
    if not target_user_email: return jsonify({"success": False, "error": "user_email required"}), 400
    from cosmos import approve_leave_request
    res = approve_leave_request(doc_id, target_user_email, admin_email, remarks)
    if res:
        html_body = f"<p>Your leave request has been <b>Approved</b>.</p><p>Remarks: {remarks}</p>"
        send_graph_email(target_user_email, "Leave Approved — Hamdaz", html_body)
    return jsonify({"success": res})
@app.route("/api/admin/leave/reject/<doc_id>", methods=["POST"])
def api_admin_reject_leave(doc_id):
    if "user" not in session: return jsonify({"success": False, "error": "Unauthorized"}), 401
    admin_email = session["user"].get("mail") or session["user"].get("userPrincipalName")
    if not is_admin(admin_email): return jsonify({"success": False, "error": "Access Denied"}), 403
    data = request.json or {}
    remarks = data.get("remarks", "")
    target_user_email = data.get("user_email")
    if not target_user_email: return jsonify({"success": False, "error": "user_email required"}), 400
    from cosmos import reject_leave_request
    res = reject_leave_request(doc_id, target_user_email, admin_email, remarks)
    if res:
        html_body = f"<p>Your leave request has been <b>Rejected</b>.</p><p>Remarks: {remarks}</p>"
        send_graph_email(target_user_email, "Leave Rejected — Hamdaz", html_body)
    return jsonify({"success": res})
# START FLASK + BACKGROUND UPDATER
# ==============================================================
threading.Thread(target=background_updater, daemon=True).start()
if __name__ == "__main__":
    app.run(debug=True)
